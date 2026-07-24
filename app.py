from flask import Flask, flash, request, render_template, redirect, url_for, jsonify, send_file, render_template_string, after_this_request, session
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import Column, Integer, String, text, BLOB, and_, not_
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape, portrait
from reportlab.lib.units import mm, cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
import os
import tempfile
from flask import send_file
# Mantenha suas importações de banco de dados (saidas, Escola, etc) e num2words
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
from werkzeug.security import generate_password_hash, check_password_hash
from openpyxl import Workbook
from openpyxl.styles import Border, Side, PatternFill, Font, Alignment, NamedStyle
from openpyxl.utils import get_column_letter, range_boundaries
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from num2words import num2words  # Para converter valores em extenso
import os
import webbrowser
from threading import Timer
from werkzeug.datastructures import ImmutableMultiDict
import json
import pandas as pd
import time  # Import time for sleep function
import os
import re
import locale
import tempfile
import pandas as pd
from sqlalchemy import func # Para busca insensível a maiúsculas/minúsculas
from unidecode import unidecode
from sqlalchemy import func
from datetime import datetime
from sqlalchemy import cast, Date
from collections import defaultdict
from sqlalchemy.sql import func, cast
from sqlalchemy.types import Date
import base64

app = Flask(__name__)
# app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///' + os.path.join(os.getcwd(), 'test.db')
# Adicione a URL do Supabase (Atenção: tem que começar com 'postgresql://')
app.config['SQLALCHEMY_DATABASE_URI'] = 'postgresql://postgres.nyuguhhlwcpuqvejhhvr:3848348349%40Irene@aws-0-ca-central-1.pooler.supabase.com:5432/postgres'
db = SQLAlchemy(app)

@app.template_filter('b64encode')
def b64encode_filter(data):
    return base64.b64encode(data).decode('utf-8')

try:
    locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')  # Tenta padrão Linux
except locale.Error:
    try:
        locale.setlocale(locale.LC_TIME, 'pt_BR')  # Tenta padrão Windows
    except locale.Error:
        pass  # Se o servidor não tiver português, ignora e não trava

#BD FORNECEDORES
class Fornecedor(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    nomeforn = db.Column(db.String(80), nullable=False)
    cnpjforn = db.Column(db.String(14), nullable=False)
    emailforn = db.Column(db.String(120), nullable=False)
    telefoneforn = db.Column(db.String(15), nullable=False)
    nomecontatoforn = db.Column(db.String(80), nullable=False)

    def __repr__(self):
        return '<Fornecedor %r>' % self.nomeforn
    
    def to_dict(self):
        return {
            'id': self.id,
            'nomeforn': self.nomeforn,
            'cnpjforn': self.cnpjforn,
            'emailforn': self.emailforn,
            'telefoneforn': self.telefoneforn,
            'nomecontatoforn': self.nomecontatoforn
        }
    
class itens(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    categoria = db.Column(db.String(90), nullable=False)
    descricaoitem = db.Column(db.String(90), nullable=False)
    unidadeitem = db.Column(db.String(120), nullable=False)

    def __repr__(self1):
        return '<itens %r>' % self1.descricaoitem
    
    def item_to_dict(self1):
        return {
            'id': self1.id,
            'categoria': self1.categoria,
            'descricaoitem': self1.descricaoitem,
            'unidadeitem': self1.unidadeitem
        }
    
class detalheitens(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    idresgatado = db.Column(db.Integer, primary_key=False)
    dataitem = db.Column(db.String, nullable=False)  # Altere o tipo de dataitem para String
    qtditem = db.Column(db.Float, nullable=False)
    valoritem = db.Column(db.Integer, nullable=False)
    fornitem = db.Column(db.String(120), nullable=False)
    idcompra = db.Column(db.Integer, primary_key=False)
    valor2 = db.Column(db.Integer, nullable=False)
    valor3 = db.Column(db.Integer, nullable=False)
    datadocumento= db.Column(db.String, nullable=False)
    numero_item = db.Column(db.Integer, nullable=False)  # Novo campo para a numeração dos itens 

    def __repr__(self2):
        return '<detalheitens %r>' % self2.idresgatado

    def to_dict(self2):
        dataitem = datetime.strptime(self2.dataitem, '%d/%m/%Y').strftime('%Y-%m-%d')  # Converta a data para o formato ISO 8601 aqui
        datadocumento = datetime.strptime(self2.datadocumento, '%d/%m/%Y').strftime('%Y-%m-%d')  # Converta a data para o formato ISO 8601 aqui
        descricaoitem = itens.query.filter_by(id=self2.idresgatado).first().descricaoitem
        return {
            'idresgatado': self2.idresgatado,
            'dataitem': dataitem,
            'qtditem': self2.qtditem,
            'valoritem': self2.valoritem,
            'fornitem': self2.fornitem,
            'valor2': self2.valor2,
            'valor3': self2.valor3,
            'datadocumento': datadocumento,
            'descricaoitem': descricaoitem,
            'numero_item': self2.numero_item
        }

    
class fontesubfonte(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    fonte = db.Column(db.String(50))
    subfonte = db.Column(db.String(50))

class entradas(db.Model):
    __tablename__ = 'entradas'
    id = db.Column(db.Integer, primary_key=True)
    data = db.Column(db.String, nullable=False)  
    capcus = db.Column(db.String, nullable=False)
    valor = db.Column(db.Integer, nullable=False)
    fonte = db.Column(db.String, nullable=False)
    subfonte = db.Column(db.String, nullable=False)
    comentario = db.Column(db.String, nullable=False)
    ano = db.Column(db.Integer, nullable=False)
    mes = db.Column(db.Integer, nullable=False)

class saidas(db.Model):
    
    __tablename__ = 'saidas'

    id = Column(Integer, primary_key=True)
    data = Column(String(10))
    fornecedor1 = Column(String(100))
    valor1 = Column(Integer)
    fornecedor2 = Column(String(100))
    valor2 = Column(Integer)
    fornecedor3 = Column(String(100))
    valor3 = Column(Integer)
    cnpj1 = Column(String(14))
    cnpj2 = Column(String(14))
    cnpj3 = Column(String(14))
    capcus = Column(String(10))
    fonte = Column(String(50))
    subfonte = Column(String(50))
    descricao = Column(String(200))
    tiponota = Column(String(10))
    numnota = Column(String(50))
    datanota = Column(String(10))
    numpag = Column(String(50))
    ano = Column(Integer)
    mes = Column(Integer)
   

    def __init__(self, data, fornecedor1, valor1, fornecedor2, valor2, fornecedor3, valor3, cnpj1, cnpj2, cnpj3, capcus, fonte, subfonte, descricao, tiponota, numnota, datanota, numpag, ano, mes):
        self.data = data
        self.fornecedor1 = fornecedor1
        self.valor1 = valor1
        self.fornecedor2 = fornecedor2
        self.valor2 = valor2
        self.fornecedor3 = fornecedor3
        self.valor3 = valor3
        self.cnpj1 = cnpj1
        self.cnpj2 = cnpj2
        self.cnpj3 = cnpj3
        self.capcus = capcus
        self.fonte = fonte
        self.subfonte = subfonte
        self.descricao = descricao
        self.tiponota = tiponota
        self.numnota = numnota
        self.datanota = datanota
        self.numpag = numpag
        self.ano = ano
        self.mes = mes

class Escola(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    nome_escola = db.Column(db.String(80), nullable=False)
    endereco = db.Column(db.String(120), nullable=False)
    cidade = db.Column(db.String(80), nullable=False)
    presidente_conselho = db.Column(db.String(80), nullable=False)
    secretario_conselho = db.Column(db.String(80), nullable=False)
    local_reuniao = db.Column(db.String(120), nullable=False)
    cnpj_conselho = db.Column(db.String(18), nullable=False)
    inep_escola = db.Column(db.String(12), nullable=False)
    nome_1_conselheiro = db.Column(db.String(80), nullable=False)
    endereco_1_conselheiro = db.Column(db.String(120), nullable=False)
    cpf_1_conselheiro = db.Column(db.String(14), nullable=False)
    nome_2_conselheiro = db.Column(db.String(80), nullable=False)
    endereco_2_conselheiro = db.Column(db.String(120), nullable=False)
    cpf_2_conselheiro = db.Column(db.String(14), nullable=False)
    nome_3_conselheiro = db.Column(db.String(80), nullable=False)
    endereco_3_conselheiro = db.Column(db.String(120), nullable=False)
    cpf_3_conselheiro = db.Column(db.String(14), nullable=False)
    nome_4_conselheiro = db.Column(db.String(80), nullable=False)
    endereco_4_conselheiro = db.Column(db.String(120), nullable=False)
    cpf_4_conselheiro = db.Column(db.String(14), nullable=False)
    ano_mandato = db.Column(db.String(4), nullable=False)

    def __repr__(self):
        return f'<Escola {self.nome_escola}>'
    
class Logos(db.Model): 
    id = db.Column(db.Integer, primary_key=True) 
    logo_escola = db.Column(db.LargeBinary, nullable=True) # Alterado aqui
    logo_estado = db.Column(db.LargeBinary, nullable=True) # Alterado aqui



#CHAMADA HOMEPAGE
@app.route('/homepage', methods=['GET'])
def homepage():
    return render_template('homepage.html')

# EXECUÇÃO DO BOTÃO DA HOME APRA ESCOLHA INSERIR OU EDITAR
@app.route('/escolher_inserir_editar_excluir', methods=['GET'])
def escolher_inserir_editar_excluir():
  # Não precisa de nenhuma lógica extra nessa rota, pois só queremos o redirecionamento
  return redirect(url_for('escolhainserir'))


@app.route('/escolhainserir', methods=['GET'])
def escolhainserir():
    return render_template('escolhainserir.html')

# INSERIR FORNECEDOR
@app.route('/inserirfornecedor', methods=['GET', 'POST'])
def inserirfornecedor():
    if request.method == 'GET':
        return render_template('InserirFornecedor.html')
    elif request.method == 'POST':
        data = request.form
        nomeforn = data['nomeforn']
        cnpjforn = data['cnpjforn']
        emailforn = data['emailforn']
        telefoneforn = data['telefoneforn']
        nomecontatoforn = data['nomecontatoforn']

        new_fornecedor = Fornecedor(nomeforn=nomeforn, cnpjforn=cnpjforn, emailforn=emailforn, telefoneforn=telefoneforn, nomecontatoforn=nomecontatoforn)

        try:
            db.session.add(new_fornecedor)
            db.session.commit()
            return {"success": True, "message": "Fornecedor inserido com sucesso!"}
        except:
            db.session.rollback()
            existing_fornecedor = Fornecedor.query.filter_by(cnpjforn=cnpjforn).first()
            flash(f'Erro ao inserir fornecedor: CNPJ já cadastrado para o fornecedor {existing_fornecedor.nomeforn}', 'error')
            return {"success": False, "message": f"CNPJ já cadastrado para o fornecedor {existing_fornecedor.nomeforn}"}


#INSERIR FORNECEDORES VIA EXCEL
@app.route('/inserir_fornecedores_excel', methods=['POST'])
def inserir_fornecedores_excel():
    if 'file' not in request.files:
        flash('Nenhum arquivo foi enviado', 'error')
        return redirect(request.url)

    file = request.files['file']
    if file.filename == '':
        flash('Nenhum arquivo selecionado', 'error')
        return redirect(request.url)

    if file and file.filename.endswith('.xlsx'):
        df = pd.read_excel(file, engine='openpyxl')

        for index, row in df.iterrows():
            nomeforn = row['nome']
            cnpjforn = row['cnpj']
            emailforn = row['email']
            telefoneforn = row['telefone']
            nomecontatoforn = row['nomecontato']

            new_fornecedor = Fornecedor(nomeforn=nomeforn, cnpjforn=cnpjforn, emailforn=emailforn, telefoneforn=telefoneforn, nomecontatoforn=nomecontatoforn)

            try:
                db.session.add(new_fornecedor)
                db.session.commit()
                time.sleep(0.02)  # Add a delay between each commit
            except:
                db.session.rollback()
                existing_fornecedor = Fornecedor.query.filter_by(cnpjforn=cnpjforn).first()
                flash(f'Erro ao inserir fornecedor: CNPJ já cadastrado para o fornecedor {existing_fornecedor.nomeforn}', 'error')

        flash('Fornecedores inseridos com sucesso!', 'success')
        return redirect(url_for('homepage'))

    else:
        flash('Tipo de arquivo inválido', 'error')
        return redirect(request.url)


#EDITAR FORNECEDOR   
@app.route('/editarfornecedor', methods=['GET', 'POST'])
def editarfornecedor():
    if request.method == 'GET':
        fornecedores = Fornecedor.query.order_by(Fornecedor.nomeforn).all()
        fornecedores = [fornecedor.to_dict() for fornecedor in fornecedores]
        return render_template('editar_excluir_fornecedor.html', fornecedores=fornecedores)
    elif request.method == 'POST':
        data = request.form
        id = data['id']
        nomeforn = data['nomeforn']
        cnpjforn = data['cnpjforn']
        emailforn = data['emailforn']
        telefoneforn = data['telefoneforn']
        nomecontatoforn = data['nomecontatoforn']

        fornecedor = Fornecedor.query.get(id)

        if 'editar' in request.form:
            fornecedor.nomeforn = nomeforn
            fornecedor.cnpjforn = cnpjforn
            fornecedor.emailforn = emailforn
            fornecedor.telefoneforn = telefoneforn 
            fornecedor.nomecontatoforn = nomecontatoforn

            try:
                db.session.commit()
                flash('Fornecedor editado com sucesso!', 'success')
            except:
                db.session.rollback()
                flash('Ocorreu um erro ao editar o fornecedor.', 'error')

        elif 'excluir' in request.form:
            try:
                db.session.delete(fornecedor)
                db.session.commit()
                flash('Fornecedor excluído com sucesso!', 'success')
            except:
                db.session.rollback()
                flash('Ocorreu um erro ao excluir o fornecedor.', 'error')

        return redirect(url_for('editarfornecedor'))

# INSERIR ITEM
@app.route('/inseriritem', methods=['GET', 'POST'])
def inseriritem():
    if request.method == 'GET':
        categorias = [row[0] for row in db.session.query(itens.categoria).distinct()]
        categorias.sort(key=lambda x: unidecode(x))
        descricaoitems_unidades = sorted(db.session.query(itens.descricaoitem, itens.unidadeitem).distinct().all(), key=lambda x: unidecode(x[0]))
        return render_template('inserir_item.html', categorias=categorias, descricaoitems_unidades=descricaoitems_unidades)
    elif request.method == 'POST':
        data = request.form
        categoria = data['categoria']
        descricaoitem = data['descricaoitem']
        unidadeitem = data['unidadeitem']

        new_item = itens(categoria=categoria, descricaoitem=descricaoitem, unidadeitem=unidadeitem)

        try:
            db.session.add(new_item)
            db.session.commit()
            flash('Item inserido com sucesso!', 'success')
        except:
            db.session.rollback()
            flash('Ocorreu um erro ao inserir o item.', 'error')

        return redirect(url_for('inseriritem'))
 
# PROCURA OS ITENS
@app.route('/get_itens', methods=['POST'])
def get_itens():
        categoria = request.form['categoria']
        Itens = itens.query.filter_by(categoria=categoria).all()
        return jsonify([item.item_to_dict() for item in Itens])

# DETALHES ITEM
@app.route('/detalhesitens', methods=['GET', 'POST'])
def detalhesitens():
    if request.method == 'GET':
        categorias = [row[0] for row in db.session.query(itens.categoria).distinct()]
        categorias.sort(key=lambda x: unidecode(x))
        return render_template('detalhes_itens.html', categorias=categorias)
    elif request.method == 'POST':
        categoria_selecionada = request.form.get('categoria')
        detalhes = db.session.query(itens.descricaoitem, itens.unidadeitem, detalheitens.dataitem, detalheitens.qtditem, detalheitens.valoritem, detalheitens.fornitem).join(itens, detalheitens.idresgatado == itens.id).filter(itens.categoria == categoria_selecionada).all()

        # Converter detalhes para dicionários
        detalhes = [{'descricaoitem': detalhe.descricaoitem, 'unidadeitem': detalhe.unidadeitem, 'dataitem': detalhe.dataitem, 'qtditem': detalhe.qtditem, 'valoritem': detalhe.valoritem, 'fornitem': detalhe.fornitem} for detalhe in detalhes]

        # Agrupar detalhes por descricaoitem
        grouped_details = defaultdict(list)
        for detalhe in detalhes:
            grouped_details[detalhe['descricaoitem']].append(detalhe)

        # Ordenar cada grupo por dataitem em ordem decrescente
        for descricaoitem, detalhes in grouped_details.items():
            detalhes.sort(key=lambda detalhe: datetime.strptime(detalhe['dataitem'], '%d/%m/%Y'), reverse=True)

        return jsonify(grouped_details)  # Retornar os detalhes como JSON 

# FONTE E SUBFONTE
@app.route('/fontes', methods=['GET', 'POST'])
def fontes():
    if request.method == 'GET':
        fontes = fontesubfonte.query.order_by(fontesubfonte.fonte).all()
        fontes_dict = defaultdict(list)
        for fonte in fontes:
            fontes_dict[fonte.fonte].append(fonte)
        fontes_dict = dict(sorted(fontes_dict.items(), key=lambda item: len(item[1])))
        max_len = max(len(subfontes) for subfontes in fontes_dict.values())
        return render_template('fontes.html', fontes=fontes_dict, max_len=max_len, len=len)
    
    elif request.method == 'POST':
        data = request.form

        if 'add' in request.form:
            fonte = data['fonte']
            subfonte = data['subfonte']
            nova_fonte_subfonte = fontesubfonte(fonte=fonte, subfonte=subfonte)
            
            try:
                db.session.add(nova_fonte_subfonte)
                db.session.commit()
                flash('Fonte/Subfonte adicionada com sucesso!', 'success')
            except:
                db.session.rollback()
                flash('Ocorreu um erro ao adicionar a Fonte/Subfonte.', 'error')

        elif 'editar' in request.form:
            id = data['id']
            fonte = data['fonte']
            subfonte = data['subfonte']

            fonte_subfonte = fontesubfonte.query.get(id)
            if fonte_subfonte:
                fonte_subfonte.fonte = fonte
                fonte_subfonte.subfonte = subfonte

                try:
                    db.session.commit()
                    flash('Fonte/Subfonte editada com sucesso!', 'success')
                except:
                    db.session.rollback()
                    flash('Ocorreu um erro ao editar a Fonte/Subfonte.', 'error')

        elif 'excluir' in request.form:
            id = data['id']
            fonte_subfonte = fontesubfonte.query.get(id)

            if fonte_subfonte:
                try:
                    db.session.delete(fonte_subfonte)
                    db.session.commit()
                    flash('Fonte/Subfonte excluída com sucesso!', 'success')
                except:
                    db.session.rollback()
                    flash('Ocorreu um erro ao excluir a Fonte/Subfonte.', 'error')

        return redirect(url_for('fontes'))


#INSERIR ENTRADA
@app.route('/inserir_entrada', methods=['GET', 'POST'])
def inserir_entrada():
    if request.method == 'GET':
        fontes = [fonte[0] for fonte in fontesubfonte.query.with_entities(fontesubfonte.fonte).distinct().all()]
        return render_template('inserir_entrada.html', fontes=fontes)
    
    elif request.method == 'POST':
        data = request.form
        data_entrada = datetime.strptime(data['data'], '%Y-%m-%d').strftime('%d/%m/%Y')
        data_entrada_1 = datetime.strptime(data['data'], '%Y-%m-%d')
        ano = data_entrada_1.year
        mes = data_entrada_1.month

        
        # Converter o valor para centavos
        valor = int(data['valor_oculto'])
        
        entrada = entradas(
            data=data_entrada,
            capcus=data['capcus'],
            valor=valor,
            fonte=data['fonte'],
            subfonte=data['subfonte'],
            comentario=data['comentario'],
            ano=ano,
            mes=mes
        )
        db.session.add(entrada)
        db.session.commit()
        return redirect(url_for('inserir_entrada'))


#PEGAR SUBFONTES
@app.route('/get_subfontes', methods=['POST'])
def get_subfontes():
    fonte = request.form['fonte']
    subfontes = fontesubfonte.query.filter_by(fonte=fonte).all()
    return jsonify([subfonte.subfonte for subfonte in subfontes])
 

#INSERIR SAÍDA
@app.route('/inserir_saida', methods=['GET', 'POST'])
def inserir_saida():
    if request.method == 'GET':
        fornecedores = [fornecedor[0] for fornecedor in Fornecedor.query.with_entities(Fornecedor.nomeforn).distinct().order_by(Fornecedor.nomeforn).all()]
        fontes = [fonte[0] for fonte in fontesubfonte.query.with_entities(fontesubfonte.fonte).distinct().order_by(fontesubfonte.fonte).all()]
        return render_template('inserir_saida.html', fornecedores=fornecedores, fontes=fontes)
    
    elif request.method == 'POST':
        data = request.form
        data_pagamento = datetime.strptime(data['data'], '%Y-%m-%d').strftime('%d/%m/%Y')
        data_pagamento_1 = datetime.strptime(data['data'], '%Y-%m-%d')
        ano = data_pagamento_1.year
        mes = data_pagamento_1.month
        
        # Converter os valores para centavos
        valor1 = int(data['valor_oculto1']) if data.get('valor_oculto1') else 0
        valor2 = int(data['valor_oculto2']) if data.get('valor_oculto2') and data['valor_oculto2'].isdigit() else 0
        valor3 = int(data['valor_oculto3']) if data.get('valor_oculto3') and data['valor_oculto3'].isdigit() else 0
        
        
        saida = saidas(
            data=data_pagamento,
            fornecedor1=data['fornecedor1'] if data.get('fornecedor1') else 0,
            valor1=valor1,
            fornecedor2=data['fornecedor2'] if data.get('fornecedor2')  else 0,
            valor2=valor2,
            fornecedor3=data['fornecedor3'] if data.get('fornecedor3')  else 0,
            valor3=valor3,
            cnpj1=Fornecedor.query.filter_by(nomeforn=data['fornecedor1']).first().cnpjforn if data.get('fornecedor1') else 0,
            cnpj2=Fornecedor.query.filter_by(nomeforn=data['fornecedor2']).first().cnpjforn if data.get('fornecedor2')  else 0,
            cnpj3=Fornecedor.query.filter_by(nomeforn=data['fornecedor3']).first().cnpjforn if data.get('fornecedor3')  else 0,
            capcus=data['capcus'],
            fonte=data['fonte'],
            subfonte=data['subfonte'],
            descricao=data['descricao'].upper(),
            tiponota=data['tiponota'],
            numnota=data['numnota'],
            datanota=datetime.strptime(data['datanota'], '%Y-%m-%d').strftime('%d/%m/%Y'),
            numpag=data['numpag'],
            ano=ano,
            mes=mes
        )
        db.session.add(saida)
        db.session.commit()
        return redirect(url_for('inserir_item_compra'))


#PEGAR CNPJ
@app.route('/get_cnpj', methods=['POST'])
def get_cnpj():
    fornecedor = request.form['fornecedor']
    cnpj = Fornecedor.query.filter_by(nomeforn=fornecedor).first().cnpjforn
    return jsonify(cnpj)



@app.route('/get_item_data', methods=['GET'])
def get_item_data():
    idcompra = request.args.get('idcompra')
    item_num = request.args.get('item_num')

    detalheitem = detalheitens.query.filter_by(idcompra=idcompra, numero_item=item_num).first()


    return jsonify(detalheitem.to_dict())



@app.route('/get_compra_detalhes/<int:idcompra>')
def get_compra_detalhes(idcompra):
    detalhes_compra = db.session.query(detalheitens, itens).\
        join(itens, detalheitens.idresgatado == itens.id).\
        filter(detalheitens.idcompra == idcompra).\
        order_by(detalheitens.numero_item).all()

    resultado = []
    for detalhe, item_info in detalhes_compra:
        resultado.append({
            'detalhe_id': detalhe.id,
            'idresgatado': detalhe.idresgatado,
            'descricaoitem': item_info.descricaoitem,
            'unidadeitem': item_info.unidadeitem,
            'numero_item': detalhe.numero_item,
            'qtditem': detalhe.qtditem,
            'valoritem': detalhe.valoritem,
            'valor2': detalhe.valor2,
            'valor3': detalhe.valor3,
            'datadocumento': detalhe.datadocumento # <-- LINHA ADICIONADA
        })
    
    return jsonify(resultado)


@app.route('/inserir_item_compra', methods=['GET', 'POST'])
def inserir_item_compra():
    if request.method == 'GET':
        itens_query = itens.query.order_by(itens.descricaoitem).all()
        itens_serializaveis = [
            {'id': item.id, 'descricaoitem': item.descricaoitem, 'unidadeitem': item.unidadeitem} 
            for item in itens_query
        ]
        categorias = [categoria[0] for categoria in itens.query.with_entities(itens.categoria).distinct().order_by(itens.categoria).all()]
        
        # 1. Pega todas as compras sem ordenar pelo banco de dados
        compras_bd = saidas.query.all()
        
        # 2. Ordena usando o Python (converte a string 'DD/MM/YYYY' para data real apenas para ordenar)
        # O reverse=True faz com que seja ordem decrescente (mais recentes primeiro)
        compras = sorted(
            compras_bd, 
            key=lambda c: datetime.strptime(c.data, '%d/%m/%Y') if c.data else datetime.min, 
            reverse=True
        )
        
        return render_template('inserir_item_compra.html', 
                               itens_=itens_serializaveis,
                               categorias=categorias, 
                               compras=compras)
    
    elif request.method == 'POST':
        data = request.form
        idcompra = data['idcompra']

        # 1. PROCESSAR EXCLUSÕES
        ids_para_deletar = request.form.getlist('itens_a_deletar[]')
        if ids_para_deletar:
            ids_validos = [int(id_str) for id_str in ids_para_deletar if id_str.isdigit()]
            if ids_validos:
                detalheitens.query.filter(detalheitens.id.in_(ids_validos)).delete(synchronize_session=False)

        # 2. PROCESSAR CRIAÇÕES E ATUALIZAÇÕES
        item_dados = {}
        sequencias = []
        
        # Agrupa os dados por ID do item
        for key, value in data.items():
            match = re.match(r'items\[(\d+)\]\[(.+)\]', key)
            if match:
                item_id, field_name = int(match.group(1)), match.group(2)
                if item_id not in item_dados:
                    item_dados[item_id] = {}
                item_dados[item_id][field_name] = value

                if field_name == 'numero_item':
                    sequencias.append(int(value))

        if len(sequencias) != len(set(sequencias)):
            flash('Erro: Existem itens com o mesmo número de sequência. Por favor, corrija.')
            return redirect(url_for('inserir_item_compra'))

        saida = saidas.query.filter_by(id=idcompra).first()
        if not saida:
            flash('Compra não encontrada!')
            return redirect(url_for('inserir_item_compra'))
        
        dataitem_base = saida.data
        fornecedor1_base = saida.fornecedor1
        # Proteção para data caso venha vazia
        try:
            datadocumento_base = datetime.strptime(data['datadocumento'], '%Y-%m-%d').strftime('%d/%m/%Y') if data.get('datadocumento') else saida.data
        except:
            datadocumento_base = saida.data
        
        for item_id, dados in item_dados.items():
            # --- PROTEÇÃO CONTRA O ERRO KEYERROR ---
            # Verifica se os campos obrigatórios estão presentes. Se não estiverem, pula este item.
            if 'numero_item' not in dados or 'qtditem_oculto' not in dados or 'valoritem_oculto' not in dados:
                print(f"Item {item_id} ignorado por dados incompletos: {dados.keys()}")
                continue
            # ---------------------------------------

            detalhe_id = dados.get('detalhe_id')
            
            # Garante que os valores numéricos sejam inteiros (padrão do banco)
            qtd = int(dados['qtditem_oculto'])
            val1 = int(dados['valoritem_oculto'])
            val2 = int(dados.get('valor2_oculto', 0) or 0)
            val3 = int(dados.get('valor3_oculto', 0) or 0)
            num_seq = int(dados['numero_item'])

            if detalhe_id:
                detalhe = detalheitens.query.get(int(detalhe_id))
                if detalhe:
                    detalhe.numero_item = num_seq
                    detalhe.qtditem = qtd
                    detalhe.valoritem = val1
                    detalhe.valor2 = val2
                    detalhe.valor3 = val3
            else:
                detalhe = detalheitens(
                    idcompra=int(idcompra),
                    idresgatado=item_id,
                    numero_item=num_seq,
                    qtditem=qtd,
                    valoritem=val1,
                    valor2=val2,
                    valor3=val3,
                    datadocumento=datadocumento_base,
                    dataitem=dataitem_base,
                    fornitem=fornecedor1_base
                )
                db.session.add(detalhe)

        db.session.commit()
        flash(f'Itens da compra {idcompra} salvos com sucesso!')
        return redirect(url_for('inserir_item_compra'))

@app.route('/cadastrar_novo_item_ajax', methods=['POST'])
def cadastrar_novo_item_ajax():
    data = request.get_json()
    try:
        novo_item = itens(
            descricaoitem=data['descricaoitem'],
            unidadeitem=data['unidadeitem'],
            categoria=data['categoria']
        )
        db.session.add(novo_item)
        db.session.commit()
        return jsonify({
            'success': True,
            'item': {
                'id': novo_item.id,
                'descricaoitem': novo_item.descricaoitem,
                'unidadeitem': novo_item.unidadeitem
            }
        }), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/get_ultimo_preco_item', methods=['GET'])
def get_ultimo_preco_item():
    iditem = request.args.get('iditem')
    if not iditem:
        return jsonify({})

    ultimo_detalhe = detalheitens.query.filter_by(idresgatado=iditem).order_by(detalheitens.id.desc()).first()

    if ultimo_detalhe:
        return jsonify({
            'valoritem': ultimo_detalhe.valoritem,
            'valor2': ultimo_detalhe.valor2,
            'valor3': ultimo_detalhe.valor3
        })
    else:
        return jsonify({})


#EXCLUIR ITEM
@app.route('/excluir_item', methods=['POST'])
def excluir_item():
    idcompra = request.form.get('idcompra')
    item_num = int(request.form.get('item_num'))
    
    # Recupera o item a ser excluído
    detalheitem = detalheitens.query.filter_by(idcompra=idcompra, numero_item=item_num).first()
    
    # Exclui o item
    db.session.delete(detalheitem)
    db.session.commit()
    
    # Atualiza a numeração dos itens restantes
    itens_restantes = detalheitens.query.filter_by(idcompra=idcompra).order_by(detalheitens.numero_item).all()
    for i, item in enumerate(itens_restantes):
        item.numero_item = i + 1
    db.session.commit()
    
    flash('Item excluído com sucesso!')
    return redirect(url_for('inserir_item_compra', idcompra=idcompra))
 
@app.route('/processar_upload_excel', methods=['POST'])
def processar_upload_excel():
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': 'Nenhum arquivo enviado.'})
    
    file = request.files['file']
    try:
        df = pd.read_excel(file, header=None)
        
        # Remove cabeçalho se existir
        if isinstance(df.iloc[0, 0], str) and 'descri' in df.iloc[0, 0].lower():
             df = df.iloc[1:]

        itens_importados = []

        # Função auxiliar de limpeza
        def limpar_valor(valor):
            try:
                if pd.isna(valor) or str(valor).strip() == '': return 0
                if isinstance(valor, str):
                    valor = valor.replace('.', '').replace(',', '.') # Ajuste conforme seu padrão BR
                return int(float(valor) * 100)
            except:
                return 0

        valores_excel = df.values.tolist()

        # O 'start=1' garante que o 'i' comece valendo 1, depois 2, depois 3...
        for i, row in enumerate(valores_excel, start=1):
            # row agora é uma lista simples: [Descricao, Qtd, Val1, Val2, Val3]
            
            descricao_excel = str(row[0]).strip() if pd.notna(row[0]) else "Item sem descrição"
            qtd = int(row[1]) if pd.notna(row[1]) else 1
            val1 = limpar_valor(row[2])
            val2 = limpar_valor(row[3])
            val3 = limpar_valor(row[4])

            itens_importados.append({
                'temp_id': f"excel_{i}", 
                'descricao_excel': descricao_excel,
                'qtditem': qtd,
                'valoritem': val1,
                'valor2': val2,
                'valor3': val3,
                'numero_item': i # Aqui garantimos que vai 1, 2, 3, 4...
            })

        return jsonify({
            'success': True,
            'itens': itens_importados
        })

    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao processar arquivo: {str(e)}'})


#TABELA PARA VER ENTRADA
@app.route('/ver_entradas')
def ver_entradas():
    return render_template('ver_entradas.html')

#FILTROS DE ENTRADA
@app.route('/get_filtros', methods=['GET'])
def get_filtros():
    anos = [entrada.data.split('/')[2] for entrada in entradas.query.with_entities(entradas.data).distinct().all()]
    fontes = [entrada.fonte for entrada in entradas.query.with_entities(entradas.fonte).distinct().all()]
    return jsonify({'anos': sorted(set(anos)), 'fontes': sorted(set(fontes))})

#PUXA AS ENTRADAS
@app.route('/get_entradas', methods=['GET'])
def get_entradas():
    ano = request.args.get('ano')
    fonte = request.args.get('fonte')
    query = entradas.query
    if ano:
        query = query.filter(entradas.data.like(f'%/{ano}'))
    if fonte:
        query = query.filter_by(fonte=fonte)

    query = query.order_by(
        text("substr(data, 7, 4) || '-' || substr(data, 4, 2) || '-' || substr(data, 1, 2)")
    )

    entradas_list = query.all()
    entradas_data = [{
        'id': entrada.id,
        'data': entrada.data,
        'capcus': entrada.capcus,
        'valor': entrada.valor,
        'fonte': entrada.fonte, 
        'subfonte': entrada.subfonte,
        'comentario': entrada.comentario
    } for entrada in entradas_list]
    return jsonify({'entradas': entradas_data})

#EDITAR ENTRADA
@app.route('/editar_entrada', methods=['GET', 'POST'])
def editar_entrada():
    if request.method == 'GET':
        entradaId = request.args.get('entradaId')
        entrada = entradas.query.get(entradaId)
        fontes = [fonte[0] for fonte in fontesubfonte.query.with_entities(fontesubfonte.fonte).distinct().all()]
        if entrada:
            entrada.data = datetime.strptime(entrada.data, '%d/%m/%Y').strftime('%Y-%m-%d')
            # Formatar o valor para exibir como R$ 4,50
            entrada.valor_formatado = f"R$ {entrada.valor / 100:.2f}".replace('.', ',')

        return render_template('editar_entrada.html', entrada=entrada, fontes=fontes)
    
    if request.method == 'POST':
        entradaId = request.args.get('entradaId')
        entrada = entradas.query.get(entradaId)
        data = request.form

        valor_oculto = data.get('valor_oculto', '').strip()
        if not valor_oculto.isdigit():
                flash('O valor informado é inválido.')
                return redirect(url_for('editar_entrada', entradaId=entradaId))
         
        # Converter o valor para centavos
        valor = int(valor_oculto)
        
        entrada.data = datetime.strptime(request.form['data'], '%Y-%m-%d').strftime('%d/%m/%Y')
        entrada.capcus = request.form['capcus']
        entrada.valor = valor
        entrada.fonte = request.form['fonte']
        entrada.subfonte = request.form['subfonte']
        entrada.comentario = request.form['comentario']
        
        db.session.commit()
        flash('Entrada atualizada com sucesso!')
        return redirect(url_for('editar_entrada', entradaId=entradaId))
    
    return render_template('editar_entrada.html', entrada=entrada)

# EXCLUIR ENTRADA
@app.route('/excluir_entrada', methods=['POST'])
def excluir_entrada():
    entrada_id = request.form.get('id') # Ou request.json.get('id') se enviar via JSON
    entrada = entradas.query.get(entrada_id)
    
    if entrada:
        try:
            db.session.delete(entrada)
            db.session.commit()
            return jsonify({'success': True, 'message': 'Entrada excluída com sucesso!'})
        except Exception as e:
            db.session.rollback()
            return jsonify({'success': False, 'message': str(e)}), 500
    
    return jsonify({'success': False, 'message': 'Entrada não encontrada.'}), 404


#TABELA PARA VER SAÍDA
@app.route('/ver_saidas')
def ver_saidas():
    return render_template('ver_saidas.html')
 
#FILTROS DE SAÍDA
@app.route('/get_filtros_saida', methods=['GET'])
def get_filtros_saida():
    anos = [saida.data.split('/')[2] for saida in saidas.query.with_entities(saidas.data).distinct().all()]
    fontes = [saida.fonte for saida in saidas.query.with_entities(saidas.fonte).distinct().all()]
    return jsonify({'anos': sorted(set(anos)), 'fontes': sorted(set(fontes))})
  
#PUXA AS SAÍDAS
@app.route('/get_saidas', methods=['GET'])
def get_saidas():
    ano = request.args.get('ano')
    fonte = request.args.get('fonte')
    query = saidas.query
    if ano:
        query = query.filter(saidas.data.like(f'%/{ano}'))
    if fonte:
        query = query.filter_by(fonte=fonte)
    
    query = query.order_by(
        text("substr(data, 7, 4) || '-' || substr(data, 4, 2) || '-' || substr(data, 1, 2)")
    )

    saidas_list = query.all()
    saidas_data = [{
        'id': saida.id,
        'data': saida.data,
        'fornecedor': saida.fornecedor1,
        'cnpj': saida.cnpj1,
        'capcus': saida.capcus,
        'descricao': saida.descricao,
        'valor': saida.valor1,
        'fonte': saida.fonte, 
        'subfonte': saida.subfonte
    } for saida in saidas_list]
    return jsonify({'saidas': saidas_data})


#EDITAR SAÍDA
@app.route('/editar_saida', methods=['GET', 'POST'])
def editar_saida():
    if request.method == 'GET':
        saidaId = request.args.get('saidaId')
        saida = saidas.query.get(saidaId)
        fornecedores = [fornecedor[0] for fornecedor in Fornecedor.query.with_entities(Fornecedor.nomeforn).distinct().order_by(Fornecedor.nomeforn).all()]
        fontes = [fonte[0] for fonte in fontesubfonte.query.with_entities(fontesubfonte.fonte).distinct().order_by(fontesubfonte.fonte).all()]
        if saida:
            saida.data = datetime.strptime(saida.data, '%d/%m/%Y').strftime('%Y-%m-%d')
            saida.datanota = datetime.strptime(saida.datanota, '%d/%m/%Y').strftime('%Y-%m-%d') if saida.datanota else ''
            # Formatar os valores para exibir como R$ 4,50
            saida.valor1 = f"R$ {saida.valor1 / 100:.2f}".replace('.', ',')
            saida.valor2 = f"R$ {saida.valor2 / 100:.2f}".replace('.', ',')
            saida.valor3 = f"R$ {saida.valor3 / 100:.2f}".replace('.', ',')
        return render_template('editar_saida.html', saida=saida, fornecedores=fornecedores, fontes=fontes)
    
    if request.method == 'POST':
        saidaId = request.args.get('saidaId')
        saida = saidas.query.get(saidaId)
        data = request.form
        data_pagamento = datetime.strptime(data['data'], '%Y-%m-%d').strftime('%d/%m/%Y')
        data_nota = datetime.strptime(data['datanota'], '%Y-%m-%d').strftime('%d/%m/%Y') if data['datanota'] else None
        cnpj1=Fornecedor.query.filter_by(nomeforn=request.form['fornecedor1']).first().cnpjforn if request.form.get('fornecedor1') else 0
        cnpj2=Fornecedor.query.filter_by(nomeforn=request.form['fornecedor2']).first().cnpjforn if request.form.get('fornecedor2')  else 0
        cnpj3=Fornecedor.query.filter_by(nomeforn=request.form['fornecedor3']).first().cnpjforn if request.form.get('fornecedor3')  else 0
        
        # Converter os valores para centavos
        valor1 = int(data['valor_oculto1'])
        valor2 = int(data['valor_oculto2'])
        valor3 = int(data['valor_oculto3'])
        
        saida.data = data_pagamento
        saida.fornecedor1 = request.form['fornecedor1']
        saida.valor1 = valor1
        saida.fornecedor2 = request.form['fornecedor2']
        saida.valor2 = valor2
        saida.fornecedor3 = request.form['fornecedor3']
        saida.valor3 = valor3
        saida.cnpj1= cnpj1
        saida.cnpj2 = cnpj2
        saida.cnpj3 = cnpj3
        saida.capcus = request.form['capcus']
        saida.fonte = request.form['fonte']
        saida.subfonte = request.form['subfonte']
        saida.descricao = request.form['descricao']
        saida.tiponota = request.form['tiponota']
        saida.numnota = request.form['numnota']
        saida.datanota = data_nota
        saida.numpag = request.form['numpag']
        
        db.session.commit()
        flash('Saída atualizada com sucesso!')
        return redirect(url_for('inserir_item_compra'))
    
    return render_template('editar_saida.html', saida=saida)

# EXCLUIR SAÍDA (E SEUS ITENS VINCULADOS)
@app.route('/excluir_saida', methods=['POST'])
def excluir_saida():
    saida_id = request.form.get('id')
    saida_obj = saidas.query.get(saida_id) # Busca na tabela 'saidas'
    
    if saida_obj:
        try:
            # 1. Primeiro exclui todos os itens vinculados a essa compra na tabela detalheitens
            # Isso busca onde 'idcompra' é igual ao id que estamos excluindo
            detalheitens.query.filter_by(idcompra=saida_id).delete()

            # 2. Depois exclui a saída (a compra em si)
            db.session.delete(saida_obj)
            
            # 3. Salva ambas as alterações no banco de uma vez
            db.session.commit()
            
            return jsonify({'success': True, 'message': 'Saída e seus itens excluídos com sucesso!'})
        except Exception as e:
            db.session.rollback() # Se der erro, desfaz tudo (não apaga nem um nem outro)
            return jsonify({'success': False, 'message': str(e)}), 500
    
    return jsonify({'success': False, 'message': 'Saída não encontrada.'}), 404 


#CONVERTER VALORES
def convert_to_float(valor):
    # Remove o prefixo 'R$ ' e os separadores de milhar
    valor = valor.replace('R$ ', '').replace('.', '')
    # Substitui a vírgula decimal por um ponto
    valor = valor.replace(',', '.')
    # Converte para float
    return float(valor)

def convert_to_string(valor):
    # Formata o valor com vírgula como separador decimal
    return f"{valor:,.2f}".replace('.', ',')

def convert_multiple_values(*valores):
    return [convert_to_string(convert_to_float(valor)) for valor in valores]

#DADOS DA ESCOLA
@app.route('/dados_escola')
def dados_escola():
    anos_mandato = sorted(set([escola.ano_mandato for escola in Escola.query.all()]))
    return render_template('dadosescola.html', anos_mandato=anos_mandato)

#CHAMA OS DADOS
@app.route('/dados/<ano_mandato>')
def get_dados(ano_mandato):
    if ano_mandato == "novo_mandato":
        dados = None
    else:
        escola = Escola.query.filter_by(ano_mandato=ano_mandato).first()
        dados = {
            'id': escola.id,
            'nome_escola': escola.nome_escola,
            'endereco': escola.endereco,
            'cidade': escola.cidade,
            'presidente_conselho': escola.presidente_conselho,
            'secretario_conselho': escola.secretario_conselho,
            'local_reuniao': escola.local_reuniao,
            'cnpj_conselho': escola.cnpj_conselho,
            'inep_escola': escola.inep_escola,
            'nome_1_conselheiro': escola.nome_1_conselheiro,
            'endereco_1_conselheiro': escola.endereco_1_conselheiro,
            'cpf_1_conselheiro': escola.cpf_1_conselheiro,
            'nome_2_conselheiro': escola.nome_2_conselheiro,
            'endereco_2_conselheiro': escola.endereco_2_conselheiro,
            'cpf_2_conselheiro': escola.cpf_2_conselheiro,
            'nome_3_conselheiro': escola.nome_3_conselheiro,
            'endereco_3_conselheiro': escola.endereco_3_conselheiro, 
            'cpf_3_conselheiro': escola.cpf_3_conselheiro,
            'nome_4_conselheiro': escola.nome_4_conselheiro,
            'endereco_4_conselheiro': escola.endereco_4_conselheiro,
            'cpf_4_conselheiro': escola.cpf_4_conselheiro,
            'ano_mandato': escola.ano_mandato
        }
    return jsonify(dados)

# SALVA OU ATUALIZA OS DADOS COM A NOVA LÓGICA
@app.route('/salvar', methods=['POST'])
def salvar():
    data = request.json
    original_ano_mandato = data.get('original_ano_mandato') # Ano original, antes de qualquer edição
    mandato_atual_ano = data.get('ano_mandato') # Ano que está no formulário

    if not mandato_atual_ano:
        return jsonify({'success': False, 'message': 'O campo "Ano do Mandato" é obrigatório.'}), 400

    try:
        # Cenário 1: ATUALIZAR um mandato existente.
        if original_ano_mandato == mandato_atual_ano and original_ano_mandato != 'novo_mandato':
            escola = Escola.query.filter_by(ano_mandato=original_ano_mandato).first()
            if escola:
                escola.nome_escola = data['nome_escola']
                escola.endereco = data['endereco']
                escola.cidade = data['cidade']
                escola.presidente_conselho = data['presidente_conselho']
                escola.secretario_conselho = data['secretario_conselho']
                escola.local_reuniao = data['local_reuniao']
                escola.cnpj_conselho = data['cnpj_conselho']
                escola.inep_escola = data['inep_escola']
                escola.nome_1_conselheiro = data['nome_1_conselheiro']
                escola.endereco_1_conselheiro = data['endereco_1_conselheiro']
                escola.cpf_1_conselheiro = data['cpf_1_conselheiro']
                escola.nome_2_conselheiro = data['nome_2_conselheiro']
                escola.endereco_2_conselheiro = data['endereco_2_conselheiro']
                escola.cpf_2_conselheiro = data['cpf_2_conselheiro']
                escola.nome_3_conselheiro = data['nome_3_conselheiro']
                escola.endereco_3_conselheiro = data['endereco_3_conselheiro']
                escola.cpf_3_conselheiro = data['cpf_3_conselheiro']
                escola.nome_4_conselheiro = data['nome_4_conselheiro']
                escola.endereco_4_conselheiro = data['endereco_4_conselheiro']
                escola.cpf_4_conselheiro = data['cpf_4_conselheiro']
                
                db.session.commit()
                return jsonify({'success': True, 'message': 'Mandato atualizado com sucesso!'})

        # Cenário 2: CRIAR um novo mandato.
        else:
            mandato_existente = Escola.query.filter_by(ano_mandato=mandato_atual_ano).first()
            if mandato_existente:
                return jsonify({'success': False, 'message': f'Já existe um mandato para o ano {mandato_atual_ano}.'}), 409

            nova_escola = Escola(
                nome_escola=data['nome_escola'],
                endereco=data['endereco'],
                cidade=data['cidade'],
                presidente_conselho=data['presidente_conselho'],
                secretario_conselho=data['secretario_conselho'],
                local_reuniao=data['local_reuniao'],
                cnpj_conselho=data['cnpj_conselho'],
                inep_escola=data['inep_escola'],
                nome_1_conselheiro=data['nome_1_conselheiro'],
                endereco_1_conselheiro=data['endereco_1_conselheiro'],
                cpf_1_conselheiro=data['cpf_1_conselheiro'],
                nome_2_conselheiro=data['nome_2_conselheiro'],
                endereco_2_conselheiro=data['endereco_2_conselheiro'],
                cpf_2_conselheiro=data['cpf_2_conselheiro'],
                nome_3_conselheiro=data['nome_3_conselheiro'],
                endereco_3_conselheiro=data['endereco_3_conselheiro'],
                cpf_3_conselheiro=data['cpf_3_conselheiro'],
                nome_4_conselheiro=data['nome_4_conselheiro'],
                endereco_4_conselheiro=data['endereco_4_conselheiro'],
                cpf_4_conselheiro=data['cpf_4_conselheiro'],
                ano_mandato=mandato_atual_ano
            )
            db.session.add(nova_escola)
            db.session.commit()
            return jsonify({'success': True, 'message': 'Novo mandato criado com sucesso!'})

    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': f'Erro ao salvar no banco: {str(e)}'}), 500

#LOGOS
@app.route('/logos') 
def logos(): 
    logos = Logos.query.first() 
    return render_template('logos.html', logos=logos)

@app.route('/salvarlogo', methods=['POST']) 
def salvarlogo(): 
    logo_escola = request.files['logo_escola'].read() if 'logo_escola' in request.files else None 
    logo_estado = request.files['logo_estado'].read() if 'logo_estado' in request.files else None 
    
    logos = Logos.query.first() 
    if logos: 
        if logo_escola:
            logos.logo_escola = logo_escola 
        if logo_estado:
            logos.logo_estado = logo_estado 
    else:
        logos = Logos(logo_escola=logo_escola, logo_estado=logo_estado)
        db.session.add(logos)
        
    db.session.commit() 
    flash('Logos salvas com sucesso!') 
    return redirect(url_for('logos')) 

@app.route('/deletarlogo', methods=['POST']) 
def deletarlogo(): 
    logos = Logos.query.first() 
    if logos: 
        db.session.delete(logos) 
        db.session.commit() 
        flash('Logos deletadas com sucesso!') 
        return redirect(url_for('logos'))


@app.route('/saldos')
def saldos():
    # Coletar os anos únicos
    anos = [entrada.data.split('/')[2] for entrada in entradas.query.with_entities(entradas.data).distinct().all()]
    anos += [saida.data.split('/')[2] for saida in saidas.query.with_entities(saidas.data).distinct().all()]
    fontes = [entrada.fonte for entrada in entradas.query.with_entities(entradas.fonte).distinct().all()]
    return render_template('saldos.html', anos=sorted(set(anos)), fontes=sorted(set(fontes)))

@app.route('/saldo', methods=['POST'])
def saldo():
    data = request.json
    ano = data.get('ano')
    fonte = data.get('fonte')
    apenas_ano = data.get('apenas_ano')
    semestre = data.get('semestre')  # 'primeiro', 'segundo' ou 'ano_todo'

    query_entradas = entradas.query
    query_saidas = saidas.query

    if fonte:
        query_entradas = query_entradas.filter_by(fonte=fonte)
        query_saidas = query_saidas.filter_by(fonte=fonte)

    entradas_lista = []
    saidas_lista = []

    if ano:
        ano_int = int(ano)

        # Primeiro: só dados até o ano escolhido
        query_entradas = query_entradas.filter(entradas.ano <= ano_int)
        query_saidas = query_saidas.filter(saidas.ano <= ano_int)

        entradas_todas = query_entradas.all()
        saidas_todas = query_saidas.all()

        if semestre == 'primeiro':
            # Anos anteriores + janeiro a junho do ano selecionado
            def filtro(item):
                try:
                    partes = item.data.split('/')
                    ano_item = int(partes[2])
                    mes_item = int(partes[1])
                    return ano_item < ano_int or (ano_item == ano_int and mes_item <= 6)
                except:
                    return False

            entradas_lista = [e for e in entradas_todas if filtro(e)]
            saidas_lista = [s for s in saidas_todas if filtro(s)]

        elif semestre == 'segundo':
            # Anos anteriores + julho a dezembro do ano selecionado
            def filtro(item):
                try:
                    partes = item.data.split('/')
                    ano_item = int(partes[2])
                    mes_item = int(partes[1])
                    return ano_item < ano_int or (ano_item == ano_int and mes_item >= 7)
                except:
                    return False

            entradas_lista = [e for e in entradas_todas if filtro(e)]
            saidas_lista = [s for s in saidas_todas if filtro(s)]

        else:  # ano_todo
            if apenas_ano:
                entradas_lista = [e for e in entradas_todas if int(e.ano) == ano_int]
                saidas_lista = [s for s in saidas_todas if int(s.ano) == ano_int]
            else:
                entradas_lista = entradas_todas
                saidas_lista = saidas_todas
    else:
        entradas_lista = query_entradas.all()
        saidas_lista = query_saidas.all()

    # Cálculos principais
    total_entradas = sum([entrada.valor for entrada in entradas_lista])
    total_saidas = sum([saida.valor1 for saida in saidas_lista])
    saldo_valor = total_entradas - total_saidas

    resultado = {
        'fonte': fonte,
        'total_entradas': total_entradas,
        'total_saidas': total_saidas,
        'saldo': saldo_valor,
        'subfontes': {},
        'custeio_consolidado': {} # Adicionado para os novos dados
    }

    # Subfontes
    subfontes = {entrada.subfonte for entrada in entradas.query.filter_by(fonte=fonte).all()}
    subfontes.update({saida.subfonte for saida in saidas.query.filter_by(fonte=fonte).all()})

    for subfonte in subfontes:
        resultado['subfontes'][subfonte] = {
            'entradas': 0,
            'saidas': 0,
            'saldo': 0
        }

    for entrada in entradas_lista:
        if entrada.subfonte in resultado['subfontes']:
            resultado['subfontes'][entrada.subfonte]['entradas'] += entrada.valor

    for saida in saidas_lista:
        if saida.subfonte in resultado['subfontes']:
            resultado['subfontes'][saida.subfonte]['saidas'] += saida.valor1

    for sub, val in resultado['subfontes'].items():
        val['saldo'] = val['entradas'] - val['saidas']

    # --- NOVO CÓDIGO: Cálculo do Custeio Consolidado ---
    total_entradas_custeio = 0
    total_saidas_custeio = 0

    for sub, valores in resultado['subfontes'].items():
        if sub != 'GERAL CAPITAL':
            total_entradas_custeio += valores['entradas']
            total_saidas_custeio += valores['saidas']
    
    saldo_custeio = total_entradas_custeio - total_saidas_custeio

    resultado['custeio_consolidado'] = {
        'entradas': total_entradas_custeio,
        'saidas': total_saidas_custeio,
        'saldo': saldo_custeio
    }
    # --- FIM DO NOVO CÓDIGO ---

    return jsonify(resultado)



# EXECUÇÃO DO BOTÃO DA HOME APRA ESCOLHA INSERIR OU EDITAR
@app.route('/abrir_documentos_compra', methods=['GET'])
def abrir_documentos_compra():
  # Não precisa de nenhuma lógica extra nessa rota, pois só queremos o redirecionamento
  return redirect(url_for('documentoscompra'))


@app.route('/documentoscompra', methods=['GET'])
def documentoscompra():
    # Obter todas as compras disponíveis
    compras = saidas.query.order_by(saidas.id.desc()).all()
    mandatos = Escola.query.order_by(Escola.ano_mandato.asc()).all()

    # 1. Pega todas as compras sem ordenar pelo banco de dados
    compras_bd = saidas.query.all()
        
        # 2. Ordena usando o Python (converte a string 'DD/MM/YYYY' para data real apenas para ordenar)
        # O reverse=True faz com que seja ordem decrescente (mais recentes primeiro)
    compras = sorted(
        compras_bd, 
            key=lambda c: datetime.strptime(c.data, '%d/%m/%Y') if c.data else datetime.min, 
            reverse=True
)
    # Recuperar a quantidade de itens já lançados para cada compra
    itens_por_compra = {}
    for compra in compras:
        itens_por_compra[compra.id] = detalheitens.query.filter_by(idcompra=compra.id).count()

    return render_template('documentoscompra.html', compras=compras, itens_por_compra=itens_por_compra, mandatos=mandatos)


from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, PageBreak
import os
import tempfile
from flask import send_file, request
from math import ceil

@app.route('/gerar_planilha', methods=['GET'])
def gerar_planilha():
    id_compra = request.args.get('idcompra')
    id_mandato = request.args.get('idmandato')
    
    if not id_compra or not id_mandato: return "IDs não fornecidos", 400

    compra = saidas.query.filter_by(id=id_compra).first()
    escola_data = Escola.query.filter_by(ano_mandato=id_mandato).first()
    itens_detalhes = detalheitens.query.filter_by(idcompra=id_compra).order_by(detalheitens.numero_item).all()

    if not compra or not escola_data: return "Dados insuficientes", 400

    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, f"consolidacao_{id_compra}.pdf")
    
    # Configuração da Página
    doc = SimpleDocTemplate(file_path, pagesize=landscape(A4), 
                            rightMargin=5*mm, leftMargin=5*mm, 
                            topMargin=2*mm, bottomMargin=2*mm)
    
    elements = []
    
    # --- 1. PREPARAR DADOS GERAIS ---
    
    # Totais Globais (Soma de tudo)
    total_a, total_b, total_c = 0, 0, 0
    ultimo_item_num = 0
    
    # Pré-calcula totais para exibir no rodapé de TODAS as páginas
    for item in itens_detalhes:
        val_total_a = (item.valoritem / 100) * item.qtditem
        val_total_b = (item.valor2 / 100) * item.qtditem
        val_total_c = (item.valor3 / 100) * item.qtditem
        total_a += val_total_a
        total_b += val_total_b
        total_c += val_total_c
        ultimo_item_num = item.numero_item

    def formatar_moeda(valor):
        return f"R$ {valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    
    def get_data_extenso(data_str):
        from datetime import datetime
        try: return datetime.strptime(data_str, '%Y-%m-%d').strftime('%d de %B de %Y')
        except:
            try: return datetime.strptime(data_str, '%d/%m/%Y').strftime('%d de %B de %Y')
            except: return data_str

    data_doc = get_data_extenso(itens_detalhes[0].datadocumento)

    # --- 2. DEFINIR ESTRUTURAS FIXAS (Cabeçalho e Rodapé) ---
    
    def criar_cabecalho():
        return [
            [compra.fonte if compra.fonte else "MAIS MERENDA", '', '', '', '', '', ''],
            ['SECRETARIA DE ESTADO DA EDUCAÇÃO', '', '', '', '', '', ''],
            ['CONSOLIDAÇÃO DE PESQUISAS DE PREÇOS', '', '', '', '', '', ''],
            # Bloco I
            ['BLOCO I - IDENTIFICAÇÃO DA UNIDADE EXECUTORA PRÓPRIA (Uex)/ENTIDADE MANTENEDORA (EM)', '', '', '', '', '', ''],
            ['01-Razão Social', '', '', '', '02-CNPJ', '', ''],
            [f"CONSELHO ESCOLAR DA {escola_data.nome_escola.upper()}", '', '', '', escola_data.cnpj_conselho, '', ''],
            # Bloco II
            ['BLOCO II - IDENTIFICAÇÃO DOS PROPONENTES (Fornecedores de produtos ou prestadores de serviços)', '', '', '', '', '', ''],
            ['03 - Razão Social do Proponente (A)', '', '', '', '04 - Razão Social do Proponente (B)', '', '05 - Razão Social do Proponente (C)'],
            [compra.fornecedor1 or "", '', '', '', compra.fornecedor2 or "", '', compra.fornecedor3 or ""],
            ['06-CNPJ do Proponente (A)', '', '', '', '07-CNPJ do Proponente (B)', '', '08-CNPJ do Proponente (C)'],
            [compra.cnpj1 or "", '', '', '', compra.cnpj2 or "", '', compra.cnpj3 or ""],
            # Bloco III Header
            ['BLOCO III - PROPOSTAS (R$ 1,00)', '', '', '', '', '', ''],
            ['09-Item', '10-Descrição dos produtos e serviços', '11-Unid.', '12-Quant.', '13-Valor proponente (A)', '14-Valor proponente (B)', '15-Valor proponente (C)']
        ] # Total 13 linhas

    def criar_rodape():
        # Bloco IV
        rows = []
        rows.append(['BLOCO IV - APURAÇÃO DAS PROPOSTAS', '', '', '', '', '', ''])
        rows.append(['16-Itens de menor valor', '', '', '', '17-Valor total dos itens de menor valor', '', ''])
        
        texto_vencedor = f"Proponente (A)\n(Itens 01 ao {ultimo_item_num})"
        rows.append([texto_vencedor, '', '', '', formatar_moeda(total_a), '', ''])
        rows.append(['Proponente (B)', '', '', '', '', '', ''])
        rows.append(['Proponente (C)', '', '', '', '', '', ''])
        rows.append(['', '', '', '', '18-Valor total', '', formatar_moeda(total_a)])
        
        # Bloco V
        rows.append(['BLOCO V - AUTENTICAÇÃO', '', '', '', '', '', ''])
        rows.append(['19-Local e data', '', '', '', '20-Nome do dirigente ou do representante legal da Uex ou da EM', '', ''])
        rows.append([f"{escola_data.cidade.upper()} - AL, {data_doc}", '', '', '', escola_data.presidente_conselho, '', ''])
        rows.append(['21-Assinatura do dirigente ou do representante legal da Uex ou da EM', '', '', '', '', '', ''])
        rows.append(['', '', '', '', '', '', '']) # Espaço assinatura
        return rows # Total 11 linhas

    # --- 3. PAGINAÇÃO MANUAL ---
    
    ITENS_POR_PAGINA = 10 # Ajuste este número se sobrar muito espaço ou cortar
    total_itens = len(itens_detalhes)
    num_paginas = ceil(total_itens / ITENS_POR_PAGINA)
    
    col_widths = [10*mm, 107*mm, 10*mm, 10*mm, 50*mm, 50*mm, 50*mm]

    for i in range(num_paginas):
        start = i * ITENS_POR_PAGINA
        end = start + ITENS_POR_PAGINA
        batch_itens = itens_detalhes[start:end]
        
# Montar dados desta página
        page_data = []
        page_data.extend(criar_cabecalho()) # Linhas 0-12
        
        # 1. Cria o estilo do parágrafo para a descrição (pode colocar antes do loop for item)
        # Importe Paragraph e ParagraphStyle no topo do seu arquivo se ainda não tiver
        style_desc = ParagraphStyle('Desc_Table', fontName='Helvetica', fontSize=7, leading=8.5)

        # Adicionar Itens
        for item in batch_itens:
            produto = itens.query.filter_by(id=item.idresgatado).first()
            desc_text = produto.descricaoitem if produto else ""
            unid = produto.unidadeitem if produto else "UN"
            
            va = (item.valoritem / 100) * item.qtditem
            vb = (item.valor2 / 100) * item.qtditem
            vc = (item.valor3 / 100) * item.qtditem
            
            # 2. Converte o texto simples em um Parágrafo para quebrar a linha automaticamente
            desc_formatada = Paragraph(desc_text, style_desc)
            
            page_data.append([
                str(item.numero_item), desc_formatada, unid, 
                str(item.qtditem).replace('.',','),
                formatar_moeda(va), formatar_moeda(vb), formatar_moeda(vc)
            ])
            
        page_data.extend(criar_rodape())
        
        # Criar Tabela
        t = Table(page_data, colWidths=col_widths)
        
        # --- ESTILIZAÇÃO DINÂMICA ---
        # Índices relativos baseados no tamanho do lote atual
        idx_item_start = 13
        idx_item_end = idx_item_start + len(batch_itens) - 1
        
        idx_ft_start = idx_item_end + 1
        
        estilo = [
            # GERAL
            ('FONTNAME', (0,0), (-1,-1), 'Helvetica'),
            ('FONTSIZE', (0,0), (-1,-1), 7), 
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('GRID', (0,0), (-1,-1), 0.5, colors.black),
            
            # 3. Mude o TOPPADDING e BOTTOMPADDING de 0 para 2. 
            # Isso garante que o texto de múltiplas linhas não encoste nas bordas.
            ('TOPPADDING', (0,0), (-1,-1), 2),
            ('BOTTOMPADDING', (0,0), (-1,-1), 2),
            
            ('LEFTPADDING', (0,0), (-1,-1), 2),
            ('RIGHTPADDING', (0,0), (-1,-1), 2),
            
            # ... (MANTENHA O RESTANTE DO SEU CÓDIGO DE ESTILO INTACTO A PARTIR DAQUI) ...
            
            # --- ESTILOS DO CABEÇALHO (Índices Fixos 0-12) ---
            ('ALIGN', (0,0), (-1,2), 'CENTER'), ('FONTNAME', (0,0), (-1,2), 'Helvetica-Bold'),
            ('SPAN', (0,0), (-1,0)), ('SPAN', (0,1), (-1,1)), ('SPAN', (0,2), (-1,2)),
            
            ('BACKGROUND', (0,3), (-1,3), colors.lightgrey), ('SPAN', (0,3), (-1,3)), ('FONTNAME', (0,3), (-1,3), 'Helvetica-Bold'),
            ('SPAN', (0,4), (3,4)), ('SPAN', (4,4), (-1,4)),
            ('SPAN', (0,5), (3,5)), ('SPAN', (4,5), (-1,5)), ('FONTNAME', (0,5), (-1,5), 'Helvetica-Bold'),
            
            ('BACKGROUND', (0,6), (-1,6), colors.lightgrey), ('SPAN', (0,6), (-1,6)), ('FONTNAME', (0,6), (-1,6), 'Helvetica-Bold'),
            ('SPAN', (0,7), (3,7)), ('SPAN', (4,7), (5,7)),
            ('SPAN', (0,8), (3,8)), ('SPAN', (4,8), (5,8)), ('FONTNAME', (0,8), (-1,8), 'Helvetica-Bold'),
            ('SPAN', (0,9), (3,9)), ('SPAN', (4,9), (5,9)),
            ('SPAN', (0,10), (3,10)), ('SPAN', (4,10), (5,10)), ('FONTNAME', (0,10), (-1,10), 'Helvetica-Bold'),
            
            ('BACKGROUND', (0,11), (-1,11), colors.lightgrey), ('SPAN', (0,11), (-1,11)), ('FONTNAME', (0,11), (-1,11), 'Helvetica-Bold'),
            ('ALIGN', (0,12), (-1,12), 'CENTER'), ('FONTNAME', (0,12), (-1,12), 'Helvetica-Bold'), ('BACKGROUND', (0,12), (-1,12), colors.whitesmoke),
            
            # --- ESTILOS DOS ITENS ---
            ('ALIGN', (0, idx_item_start), (0, idx_item_end), 'CENTER'),
            ('ALIGN', (1, idx_item_start), (1, idx_item_end), 'LEFT'),
            ('ALIGN', (2, idx_item_start), (3, idx_item_end), 'CENTER'),
            ('ALIGN', (4, idx_item_start), (-1, idx_item_end), 'RIGHT'),

            # --- ESTILOS DO RODAPÉ (Calculados) ---
            # Bloco IV Título
            ('BACKGROUND', (0, idx_ft_start), (-1, idx_ft_start), colors.lightgrey),
            ('SPAN', (0, idx_ft_start), (-1, idx_ft_start)),
            ('FONTNAME', (0, idx_ft_start), (-1, idx_ft_start), 'Helvetica-Bold'),
            
            # Sub-headers Bloco IV
            ('SPAN', (0, idx_ft_start+1), (3, idx_ft_start+1)),
            ('SPAN', (4, idx_ft_start+1), (-1, idx_ft_start+1)),
            ('FONTNAME', (0, idx_ft_start+1), (-1, idx_ft_start+1), 'Helvetica-Bold'),
            
            # Proponente A, B, C
            ('SPAN', (0, idx_ft_start+2), (3, idx_ft_start+2)), ('SPAN', (4, idx_ft_start+2), (-1, idx_ft_start+2)),
            ('SPAN', (0, idx_ft_start+3), (3, idx_ft_start+3)), ('SPAN', (4, idx_ft_start+3), (-1, idx_ft_start+3)),
            ('SPAN', (0, idx_ft_start+4), (3, idx_ft_start+4)), ('SPAN', (4, idx_ft_start+4), (-1, idx_ft_start+4)),
            
            # Total 18
            ('SPAN', (4, idx_ft_start+5), (5, idx_ft_start+5)),
            ('ALIGN', (4, idx_ft_start+5), (-1, idx_ft_start+5), 'RIGHT'),
            ('FONTNAME', (0, idx_ft_start+5), (-1, idx_ft_start+5), 'Helvetica-Bold'),
            
            # Bloco V Título
            ('BACKGROUND', (0, idx_ft_start+6), (-1, idx_ft_start+6), colors.lightgrey),
            ('SPAN', (0, idx_ft_start+6), (-1, idx_ft_start+6)),
            ('FONTNAME', (0, idx_ft_start+6), (-1, idx_ft_start+6), 'Helvetica-Bold'),
            
            # Local/Data Labels
            ('SPAN', (0, idx_ft_start+7), (3, idx_ft_start+7)),
            ('SPAN', (4, idx_ft_start+7), (-1, idx_ft_start+7)),
            
            # Local/Data Dados
            ('SPAN', (0, idx_ft_start+8), (3, idx_ft_start+8)),
            ('SPAN', (4, idx_ft_start+8), (-1, idx_ft_start+8)),
            ('FONTNAME', (0, idx_ft_start+8), (-1, idx_ft_start+8), 'Helvetica-Bold'),
            
            # Assinatura Header (21-Assinatura...)
            ('SPAN', (0, idx_ft_start+9), (-1, idx_ft_start+9)),
            # Removendo borda inferior desta linha para conectar com a caixa de assinatura
            ('LINEBELOW', (0, idx_ft_start+9), (-1, idx_ft_start+9), 0, colors.white),
            
            # Caixa Assinatura (Vazia)
            ('SPAN', (0, idx_ft_start+10), (-1, idx_ft_start+10)),
            ('MINROWHEIGHT', (0, idx_ft_start+10), (-1, idx_ft_start+10), 15*mm),
            # Removendo borda superior (para conectar)
            ('LINEABOVE', (0, idx_ft_start+10), (-1, idx_ft_start+10), 0, colors.white),
            # Removendo bordas internas (já garantido pelo SPAN -1, mas reforçando visualmente)
            ('BOX', (0, idx_ft_start+9), (-1, idx_ft_start+10), 0.5, colors.black),
        ]
        
        t.setStyle(TableStyle(estilo))
        elements.append(t)
        
        # Quebra de página, exceto na última
        if i < num_paginas - 1:
            elements.append(PageBreak())

    doc.build(elements)
    return send_file(file_path, as_attachment=True, download_name=f"consolidacao_{id_compra}.pdf")


@app.route('/gerar_ordem', methods=['GET'])
def gerar_ordem():
    id_compra = request.args.get('idcompra')
    id_mandato = request.args.get('idmandato')
    
    if not id_compra or not id_mandato: return "Erro IDs", 400

    compra = saidas.query.filter_by(id=id_compra).first()
    escola_data = Escola.query.filter_by(ano_mandato=id_mandato).first()
    itens_detalhes = detalheitens.query.filter_by(idcompra=id_compra).order_by(detalheitens.numero_item).all()

    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, f"ordem_{id_compra}.pdf")
    
    # Ordem de compra geralmente é Retrato (Portrait)
    doc = SimpleDocTemplate(file_path, pagesize=portrait(A4), leftMargin=2*mm, rightMargin=2*mm, topMargin=2*mm, bottomMargin=2*mm)
    elements = []
    
    # --- CABEÇALHO ---
    # Criamos uma tabela para o cabeçalho para organizar logo/textos
    header_data = [
        ['', 'ORDEM DE COMPRA/SERVIÇO', ' '],
        ['', compra.fonte, ''],
        ['UEX CONTRATANTE:', f"CONSELHO ESCOLAR DA {escola_data.nome_escola}", 'CNPJ: ' + escola_data.cnpj_conselho],
        ['NOME DA ESCOLA:', escola_data.nome_escola, 'SEEC: ' + (escola_data.inep_escola or "")],
        ['PROPONENTE VENCEDOR:', compra.fornecedor1, 'CNPJ: ' + (compra.cnpj1 or "")],
    ]
    
    t_head = Table(header_data, colWidths=[50*mm, 80*mm, 60*mm])
    t_head.setStyle(TableStyle([
        ('FONTNAME', (0,0), (-1,-1), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (1,-1), 9),
        ('ALIGN',(0,0),(-1,-1),'CENTER'),
        ('FONTSIZE', (0,2), (-1,-1), 6.5),
        ('GRID', (0,0), (-1,-1), 0.5, colors.black),
        ('BACKGROUND', (0,0), (-1,0), colors.lightgrey), # Cor de fundo no titulo
        ('SPAN', (2,0), (-1,0)), # Mescla titulo direita
    ]))
    elements.append(t_head)
    elements.append(Spacer(1, 3*mm))

    # Texto Legal
    p_style = ParagraphStyle('Justify', alignment=TA_JUSTIFY, fontSize=9, fontName='Helvetica')
    texto_lei = "Autorizo o fornecimento dos itens, conforme descrição na planilha abaixo, em razão do proponente acima identificado ter apresentado uma proposta adequada e de menor preço."
    elements.append(Paragraph(texto_lei, p_style))
    elements.append(Spacer(1, 3*mm))

# --- TABELA DE ITENS ---
    items_data = [['ITEM', 'DESCRIÇÃO', 'UNID', 'QUANT', 'V. UNIT', 'V. TOTAL']]
    
    total_geral = 0
    def fmt(v): return f"R$ {v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

    # 1. Criamos um estilo específico para a descrição quebrar a linha dentro da tabela
    style_desc = ParagraphStyle('Desc_Table', fontName='Helvetica', fontSize=7, leading=8.5)

    for idx, item in enumerate(itens_detalhes, 1):
        desc = itens.query.filter_by(id=item.idresgatado).first().descricaoitem
        unid = itens.query.filter_by(id=item.idresgatado).first().unidadeitem
        v_unit = item.valoritem / 100
        v_total = v_unit * item.qtditem
        total_geral += v_total
        
        # 2. Envolvemos a descrição (desc) no objeto Paragraph usando o estilo que criamos
        desc_formatada = Paragraph(desc, style_desc)
        
        # 3. Adicionamos a desc_formatada em vez da string comum
        items_data.append([str(idx), desc_formatada, unid, str(item.qtditem), fmt(v_unit), fmt(v_total)])

    # Linha Total
    items_data.append(['', '', '', '', 'TOTAL:', fmt(total_geral)])

    t_items = Table(items_data, colWidths=[10*mm, 80*mm, 15*mm, 15*mm, 30*mm, 30*mm], repeatRows=1)
    
    # Estilo Tabela Itens
    style_items = [
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'), # Header negrito
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#C4BD97")), # Cor do Header (igual planilha)
        ('GRID', (0,0), (-1,-2), 0.5, colors.black), # Bordas nos itens
        ('ALIGN', (3,0), (-1,-1), 'RIGHT'), # Valores alinhados a direita
        ('ALIGN', (0,0), (2,-1), 'LEFT'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'), # 4. Adicionado VALIGN MIDDLE para os números não ficarem no topo quando a linha quebrar
        ('FONTSIZE', (0,0), (-1,-1), 7),
        # Linha Total
        ('FONTNAME', (-2,-1), (-1,-1), 'Helvetica-Bold'),
        ('BACKGROUND', (-2,-1), (-1,-1), colors.HexColor("#C4BD97")),
        ('GRID', (-2,-1), (-1,-1), 0.5, colors.black),
        ('TOPPADDING', (0, 0), (-1, -1), 2),    # 5. Aumentado o padding para 2 para o texto não colar nas bordas quando tiver 2 linhas
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2), 
    ]
    t_items.setStyle(TableStyle(style_items))
    elements.append(t_items)
    
    # --- RODAPÉ / ASSINATURAS ---
    elements.append(Spacer(1, 10*mm))
    
    from datetime import datetime
    dt_fmt = datetime.strptime(itens_detalhes[0].datadocumento, '%d/%m/%Y').strftime('%d de %B de %Y')
    
    # Tabela invisivel para assinaturas lado a lado
    sig_data = [
        ['LOCAL E DATA', 'PRESIDENTE DO CONSELHO'],
        [f"{escola_data.cidade.upper()} - AL\n{dt_fmt}", ''],
        ['', '__________________________________'],
        ['', escola_data.presidente_conselho]
    ]
    t_sig = Table(sig_data, colWidths=[90*mm, 90*mm])
    t_sig.setStyle(TableStyle([
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
    ]))
    elements.append(t_sig)

    doc.build(elements)
    return send_file(file_path, as_attachment=True, download_name=f"ordem_{id_compra}.pdf")


def formatar_valor(valor):
    return f"R$ {valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def valor_por_extenso(valor):
    reais = int(valor)
    centavos = int(round((valor - reais) * 100))
    if centavos > 0:
        return f"{num2words(reais, lang='pt_BR')} reais e {num2words(centavos, lang='pt_BR')} centavos".upper()
    else:
        return f"{num2words(reais, lang='pt_BR')} reais".upper()

@app.route('/gerar_ata', methods=['GET'])
def gerar_ata():
    id_compra = request.args.get('idcompra')
    id_mandato = request.args.get('idmandato')
    if not id_compra or not id_mandato: return "Erro", 400

    # 1. BUSCAR DADOS (Faltava a busca dos itens_detalhes aqui)
    logos = Logos.query.first()
    compra = saidas.query.filter_by(id=id_compra).first()
    escola_data = Escola.query.filter_by(ano_mandato=id_mandato).first()
    
    # ESTA É A LINHA QUE FALTAVA:
    itens_detalhes = detalheitens.query.filter_by(idcompra=id_compra).order_by(detalheitens.numero_item).all()
    
    if not itens_detalhes:
        return "Nenhum item encontrado para esta compra", 400

    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, f"ata_{id_compra}.pdf")
    
    doc = SimpleDocTemplate(file_path, pagesize=portrait(A4), leftMargin=5*mm, rightMargin=5*mm, topMargin=7*mm, bottomMargin=2*mm, fontSize=9)
    elements = []
    styles = getSampleStyleSheet()
    
    # --- LOGO ---
    logo_path = None
    if logos:
        if logos.logo_escola:
            logo_path = os.path.join(temp_dir, "logo_tmp.png")
            with open(logo_path, "wb") as f: f.write(logos.logo_escola)
        elif logos.logo_estado:
            logo_path = os.path.join(temp_dir, "logo_tmp.png")
            with open(logo_path, "wb") as f: f.write(logos.logo_estado)
            
    if logo_path:
        im = RLImage(logo_path, width=40*mm, height=40*mm, kind='proportional')
        im.hAlign = 'CENTER'
        elements.append(im)
        elements.append(Spacer(1, 5*mm))

    # --- TÍTULO ---
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], alignment=TA_CENTER, fontSize=12)
    elements.append(Paragraph(compra.descricao, title_style))
    elements.append(Spacer(1, 5*mm))
    
    # --- LISTA PROPOSTAS (Cabeçalho da Ata) ---
    normal_style = ParagraphStyle('Normal_Custom', parent=styles['Normal'], fontSize=10, leading=12)
    
    def fmt(v): return f"R$ {v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    
    propostas = [
        {"nome": compra.fornecedor1, "cnpj": compra.cnpj1, "valor": compra.valor1/100},
        {"nome": compra.fornecedor2, "cnpj": compra.cnpj2, "valor": compra.valor2/100},
        {"nome": compra.fornecedor3, "cnpj": compra.cnpj3, "valor": compra.valor3/100},
    ]

    for idx, p in enumerate(propostas, 1):
        if p['nome']:
            # Usando num2words para extenso, se tiver importado
            v_extenso = valor_por_extenso(p['valor']) 
            txt = f"<b>{idx}- {p['nome']}</b><br/>CNPJ: {p['cnpj']}<br/>Valor: {fmt(p['valor'])} <br/> ({v_extenso})"
            elements.append(Paragraph(txt, normal_style))
            elements.append(Spacer(1, 3*mm))

    elements.append(PageBreak()) 

    # --- TEXTO DA ATA ---
    elements.append(Paragraph(f"CONSELHO ESCOLAR DA {escola_data.nome_escola.upper()}", title_style))
    elements.append(Paragraph("ATA DA ANÁLISE E HOMOLOGAÇÃO", title_style))
    elements.append(Spacer(1, 7*mm))
    
    from datetime import datetime
    # Agora itens_detalhes existe e não dará erro
    dt_fmt = datetime.strptime(itens_detalhes[0].datadocumento, '%d/%m/%Y').strftime('%d de %B de %Y')

    texto_corpo = f"""
    Aos {dt_fmt}, reuniram-se os membros do Conselho Escolar da {escola_data.nome_escola}, situada na {escola_data.endereco}, 
    na {escola_data.local_reuniao}. O presidente {escola_data.presidente_conselho}, designou-me para secretariar esta ata, 
    e apresentou as propostas para a aquisição dos bens ofertados pelo programa {compra.capcus} já pesquisados, 
    conforme entendimento em reunião anterior, com recursos oriundos do {compra.fonte}, entre as empresas:
    """
    justify_style = ParagraphStyle('Justify', parent=styles['Normal'], alignment=TA_JUSTIFY, fontSize=11, leading=14)
    elements.append(Paragraph(texto_corpo, justify_style))
    elements.append(Spacer(1, 5*mm))

    # Repete propostas no corpo
    for idx, p in enumerate(propostas, 1):
        if p['nome']:
            v_extenso = valor_por_extenso(p['valor'])
            elements.append(Paragraph(f"{idx}- {p['nome']} <br/>CNPJ: {p['cnpj']}<br/>Valor: {fmt(p['valor'])}<br/>({v_extenso})", normal_style))
            elements.append(Spacer(1, 1*mm))
            

    texto_final = f"""
    <br/><br/>Portanto, verificamos que a empresa <b>{compra.fornecedor1}</b> apresentou a melhor proposta. 
    Assim, autorizamos o fornecimento dos bens/serviços solicitados ofertados pelo programa {compra.capcus} solicitados, atendendo as normas do {compra.fonte}. Nada mais havendo a tratar, deu-se por encerrada a reunião e esta ata será lavrada e assinada por mim, {escola_data.secretario_conselho} e pelos demais presentes.
    """
    elements.append(Paragraph(texto_final, justify_style))
    elements.append(Spacer(1, 20*mm))

    # --- ASSINATURAS ---
    sig_lines = [['ASSINATURA', 'SEGMENTO']]
    for _ in range(8):
        sig_lines.append(['___________________________________________', '_______________________'])
        
    t_sigs = Table(sig_lines, colWidths=[100*mm, 60*mm])
    t_sigs.setStyle(TableStyle([
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'BOTTOM'),
        ('FONTSIZE', (0,0), (-1,-1), 10),
        ('BOTTOMPADDING', (0,0), (-1,-1), 8*mm),
    ]))
    elements.append(t_sigs)

    doc.build(elements)
    return send_file(file_path, as_attachment=True, download_name=f"ata_{id_compra}.pdf")

@app.route('/abrir_documentos_relatorio', methods=['GET'])
def abrir_documentos_relatorio():
  # Não precisa de nenhuma lógica extra nessa rota, pois só queremos o redirecionamento
  return redirect(url_for('documentosrelatorio'))


@app.route('/documentosrelatorio', methods=['GET'])
def documentosrelatorio():
    # Obter todas as compras disponíveis
    anos = [entrada.data.split('/')[2] for entrada in entradas.query.with_entities(entradas.data).distinct().all()]
    anos += [saida.data.split('/')[2] for saida in saidas.query.with_entities(saidas.data).distinct().all()]
    fontes = [entrada.fonte for entrada in entradas.query.with_entities(entradas.fonte).distinct().all()]
    mandato = [dados.ano_mandato for dados in Escola.query.with_entities(Escola.ano_mandato).distinct().all()]
   

    return render_template('relatorios.html', anos=sorted(set(anos)), fontes=sorted(set(fontes)), mandato=sorted(set(mandato)))

 # Rota para gerar o documento da Ata


@app.route('/gerar_parecer', methods=['GET'])
def gerar_parecer():
    ano_dados = request.args.get('ano')
    fonte_dados=request.args.get('fonte')
    mandato_dados = request.args.get('anoMandato')
    data_dados = request.args.get('data')
    mes_dados = request.args.get('mes')

    escola_data = Escola.query.filter_by(ano_mandato=mandato_dados).first()


    query_entradas = entradas.query
    query_saidas = saidas.query

    query_entradas = query_entradas.filter_by(fonte=fonte_dados)
    query_saidas = query_saidas.filter_by(fonte=fonte_dados)

    
    ano_int = int(ano_dados)
    ano_anterior = ano_int-1
    periodo = int(mes_dados)

    if periodo == 3:
        dadoperiodo = ano_dados
        mesinicial = "01"
        mesfinal = "12"
        diafinal = "31"
        #Saldo do exercício anterior
        query_entradas_anterior_custeio = query_entradas.filter(entradas.ano <= ano_anterior, entradas.capcus.like("CUSTEIO"))
        query_saidas_anterior_custeio = query_saidas.filter(saidas.ano <= ano_anterior, saidas.capcus.like("CUSTEIO"))
        query_entradas_anterior_capital = query_entradas.filter(entradas.ano <= ano_anterior, entradas.capcus.like("CAPITAL")) 
        query_saidas_anterior_capital = query_saidas.filter(saidas.ano <= ano_anterior, saidas.capcus.like("CAPITAL"))    

        entradas_lista_anterior_custeio = query_entradas_anterior_custeio.all()
        entradas_lista_anterior_capital = query_entradas_anterior_capital.all()
        saidas_lista_anterior_custeio = query_saidas_anterior_custeio.all()
        saidas_lista_anterior_capital = query_saidas_anterior_capital.all()
        
    
        total_entradas_anterior_custeio = sum([entrada.valor for entrada in entradas_lista_anterior_custeio])
        total_entradas_anterior_capital = sum([entrada.valor for entrada in entradas_lista_anterior_capital])
        total_saidas_anterior_custeio = sum([saida.valor1 for saida in saidas_lista_anterior_custeio])
        total_saidas_anterior_capital = sum([saida.valor1 for saida in saidas_lista_anterior_capital])

        saldo_anterior_custeio = total_entradas_anterior_custeio - total_saidas_anterior_custeio
        saldo_anterior_capital = total_entradas_anterior_capital - total_saidas_anterior_capital

        saldo_anterior_custeio = saldo_anterior_custeio/100
        saldo_anterior_capital = saldo_anterior_capital/100
        
        #Recurso financeiro recebido
        query_entradas_recebido_custeio = query_entradas.filter(entradas.ano == ano_int,  entradas.capcus.like("CUSTEIO"))
        query_entradas_recebido_capital = query_entradas.filter(entradas.ano == ano_int, entradas.capcus.like("CAPITAL"))
        
        entradas_lista_recebido_custeio = query_entradas_recebido_custeio.all()
        entradas_lista_recebido_capital = query_entradas_recebido_capital.all()

        total_entradas_recebido_custeio = sum([entrada.valor for entrada in entradas_lista_recebido_custeio])
        total_entradas_recebido_capital = sum([entrada.valor for entrada in entradas_lista_recebido_capital])

        total_entradas_recebido_custeio = total_entradas_recebido_custeio/100
        total_entradas_recebido_capital = total_entradas_recebido_capital/100
        
        #Recursos próprios
        query_entradas_proprio_custeio = query_entradas.filter(entradas.ano == ano_int, entradas.comentario == "RECURSOS PRÓPRIOS", entradas.capcus.like("CUSTEIO"))
        query_entradas_proprio_capital = query_entradas.filter(entradas.ano == ano_int, entradas.comentario == "RECURSOS PRÓPRIOS", entradas.capcus.like("CAPITAL"))
        
        entradas_lista_proprio_custeio = query_entradas_proprio_custeio.all()
        entradas_lista_proprio_capital = query_entradas_proprio_capital.all()

        total_entradas_proprio_custeio = sum([entrada.valor for entrada in entradas_lista_proprio_custeio])
        total_entradas_proprio_capital = sum([entrada.valor for entrada in entradas_lista_proprio_capital])

        total_entradas_proprio_custeio = total_entradas_proprio_custeio/100
        total_entradas_proprio_capital = total_entradas_proprio_capital/100

        #Rendimentos
        query_entradas_rendimento_custeio = query_entradas.filter(entradas.ano == ano_int, entradas.comentario == "RENDIMENTOS", entradas.capcus.like("CUSTEIO"))
        query_entradas_rendimento_capital = query_entradas.filter(entradas.ano == ano_int, entradas.comentario == "RENDIMENTOS", entradas.capcus.like("CAPITAL"))
        
        entradas_lista_rendimento_custeio = query_entradas_rendimento_custeio.all()
        entradas_lista_rendimento_capital = query_entradas_rendimento_capital.all()

        total_entradas_rendimento_custeio = sum([entrada.valor for entrada in entradas_lista_rendimento_custeio])
        total_entradas_rendimento_capital = sum([entrada.valor for entrada in entradas_lista_rendimento_capital])

        total_entradas_rendimento_custeio = total_entradas_rendimento_custeio/100
        total_entradas_rendimento_capital = total_entradas_rendimento_capital/100

        #Recurso calculado


        total_entradas_recebido_custeio = total_entradas_recebido_custeio - total_entradas_proprio_custeio - total_entradas_rendimento_custeio
        total_entradas_recebido_capital = total_entradas_recebido_capital - total_entradas_proprio_capital - total_entradas_rendimento_capital

        #Receita total

        total_receita_custeio = saldo_anterior_custeio + total_entradas_recebido_custeio + total_entradas_proprio_custeio + total_entradas_rendimento_custeio

        total_receita_capital = saldo_anterior_capital + total_entradas_recebido_capital + total_entradas_proprio_capital + total_entradas_rendimento_capital

        #Despesas

        query_saidas_custeio = query_saidas.filter(saidas.ano == ano_int, saidas.capcus.like("CUSTEIO"))
        query_saidas_capital = query_saidas.filter(saidas.ano == ano_int, saidas.capcus.like("CAPITAL"))

        saidas_lista_custeio = query_saidas_custeio.all()
        saidas_lista_capital = query_saidas_capital.all()

        total_saidas_custeio = sum([saida.valor1 for saida in saidas_lista_custeio])
        total_saidas_capital = sum([saida.valor1 for saida in saidas_lista_capital])

        total_saidas_custeio = total_saidas_custeio/100
        total_saidas_capital = total_saidas_capital/100

        #Saldo final

        saldo_final_custeio = total_receita_custeio - total_saidas_custeio
        saldo_final_capital = total_receita_capital - total_saidas_capital
    
    if periodo == 1:
        dadoperiodo = f'{ano_dados}.1'
        mesinicial = "01"
        mesfinal = "06"
        diafinal = "30"
        #Saldo do exercício anterior
        query_entradas_anterior_custeio = query_entradas.filter(entradas.ano <= ano_anterior, entradas.capcus.like("CUSTEIO"))
        query_saidas_anterior_custeio = query_saidas.filter(saidas.ano <= ano_anterior, saidas.capcus.like("CUSTEIO"))
        query_entradas_anterior_capital = query_entradas.filter(entradas.ano <= ano_anterior, entradas.capcus.like("CAPITAL")) 
        query_saidas_anterior_capital = query_saidas.filter(saidas.ano <= ano_anterior, saidas.capcus.like("CAPITAL"))    

        entradas_lista_anterior_custeio = query_entradas_anterior_custeio.all()
        entradas_lista_anterior_capital = query_entradas_anterior_capital.all()
        saidas_lista_anterior_custeio = query_saidas_anterior_custeio.all()
        saidas_lista_anterior_capital = query_saidas_anterior_capital.all()
    
        total_entradas_anterior_custeio = sum([entrada.valor for entrada in entradas_lista_anterior_custeio])
        total_entradas_anterior_capital = sum([entrada.valor for entrada in entradas_lista_anterior_capital])
        total_saidas_anterior_custeio = sum([saida.valor1 for saida in saidas_lista_anterior_custeio])
        total_saidas_anterior_capital = sum([saida.valor1 for saida in saidas_lista_anterior_capital])

        saldo_anterior_custeio = total_entradas_anterior_custeio - total_saidas_anterior_custeio
        saldo_anterior_capital = total_entradas_anterior_capital - total_saidas_anterior_capital

        saldo_anterior_custeio = saldo_anterior_custeio/100
        saldo_anterior_capital = saldo_anterior_capital/100
        
        #Recurso financeiro recebido
        query_entradas_recebido_custeio = query_entradas.filter(entradas.ano == ano_int, entradas.mes <= 6,  entradas.capcus.like("CUSTEIO"))
        query_entradas_recebido_capital = query_entradas.filter(entradas.ano == ano_int, entradas.mes <= 6, entradas.capcus.like("CAPITAL"))
        
        entradas_lista_recebido_custeio = query_entradas_recebido_custeio.all()
        entradas_lista_recebido_capital = query_entradas_recebido_capital.all()

        total_entradas_recebido_custeio = sum([entrada.valor for entrada in entradas_lista_recebido_custeio])
        total_entradas_recebido_capital = sum([entrada.valor for entrada in entradas_lista_recebido_capital])

        total_entradas_recebido_custeio = total_entradas_recebido_custeio/100
        total_entradas_recebido_capital = total_entradas_recebido_capital/100
        
        #Recursos próprios
        query_entradas_proprio_custeio = query_entradas.filter(entradas.ano == ano_int, entradas.mes <= 6, entradas.comentario == "RECURSOS PRÓPRIOS", entradas.capcus.like("CUSTEIO"))
        query_entradas_proprio_capital = query_entradas.filter(entradas.ano == ano_int, entradas.mes <= 6, entradas.comentario == "RECURSOS PRÓPRIOS", entradas.capcus.like("CAPITAL"))
        
        entradas_lista_proprio_custeio = query_entradas_proprio_custeio.all()
        entradas_lista_proprio_capital = query_entradas_proprio_capital.all()

        total_entradas_proprio_custeio = sum([entrada.valor for entrada in entradas_lista_proprio_custeio])
        total_entradas_proprio_capital = sum([entrada.valor for entrada in entradas_lista_proprio_capital])

        total_entradas_proprio_custeio = total_entradas_proprio_custeio/100
        total_entradas_proprio_capital = total_entradas_proprio_capital/100

        #Rendimentos
        query_entradas_rendimento_custeio = query_entradas.filter(entradas.ano == ano_int, entradas.mes <= 6, entradas.comentario == "RENDIMENTOS", entradas.capcus.like("CUSTEIO"))
        query_entradas_rendimento_capital = query_entradas.filter(entradas.ano == ano_int, entradas.mes <= 6, entradas.comentario == "RENDIMENTOS", entradas.capcus.like("CAPITAL"))
        
        entradas_lista_rendimento_custeio = query_entradas_rendimento_custeio.all()
        entradas_lista_rendimento_capital = query_entradas_rendimento_capital.all()

        total_entradas_rendimento_custeio = sum([entrada.valor for entrada in entradas_lista_rendimento_custeio])
        total_entradas_rendimento_capital = sum([entrada.valor for entrada in entradas_lista_rendimento_capital])

        total_entradas_rendimento_custeio = total_entradas_rendimento_custeio/100
        total_entradas_rendimento_capital = total_entradas_rendimento_capital/100

        #Recurso calculado


        total_entradas_recebido_custeio = total_entradas_recebido_custeio - total_entradas_proprio_custeio - total_entradas_rendimento_custeio
        total_entradas_recebido_capital = total_entradas_recebido_capital - total_entradas_proprio_capital - total_entradas_rendimento_capital

        #Receita total

        total_receita_custeio = saldo_anterior_custeio + total_entradas_recebido_custeio + total_entradas_proprio_custeio + total_entradas_rendimento_custeio

        total_receita_capital = saldo_anterior_capital + total_entradas_recebido_capital + total_entradas_proprio_capital + total_entradas_rendimento_capital

        #Despesas

        query_saidas_custeio = query_saidas.filter(saidas.ano == ano_int, saidas.mes <= 6, saidas.capcus.like("CUSTEIO"))
        query_saidas_capital = query_saidas.filter(saidas.ano == ano_int, saidas.mes <= 6, saidas.capcus.like("CAPITAL"))

        saidas_lista_custeio = query_saidas_custeio.all()
        saidas_lista_capital = query_saidas_capital.all()

        total_saidas_custeio = sum([saida.valor1 for saida in saidas_lista_custeio])
        total_saidas_capital = sum([saida.valor1 for saida in saidas_lista_capital])

        total_saidas_custeio = total_saidas_custeio/100
        total_saidas_capital = total_saidas_capital/100

        #Saldo final

        saldo_final_custeio = total_receita_custeio - total_saidas_custeio
        saldo_final_capital = total_receita_capital - total_saidas_capital
    
    if periodo == 2:
        dadoperiodo = f'{ano_dados}.2'
        mesinicial = "06"
        mesfinal = "12"
        diafinal = "31"
        #Saldo do exercício anterior
        query_entradas_anterior_custeio = query_entradas.filter(entradas.ano == ano_int, entradas.mes <= 6, entradas.capcus.like("CUSTEIO"))
        query_saidas_anterior_custeio = query_saidas.filter(saidas.ano == ano_int, saidas.mes <= 6, saidas.capcus.like("CUSTEIO"))
        query_entradas_anterior_capital = query_entradas.filter(entradas.ano == ano_int, entradas.mes <= 6, entradas.capcus.like("CAPITAL")) 
        query_saidas_anterior_capital = query_saidas.filter(saidas.ano == ano_int, saidas.mes <= 6, saidas.capcus.like("CAPITAL"))

        query_entradas_anterior_custeio_ = query_entradas.filter(entradas.ano <= ano_anterior, entradas.capcus.like("CUSTEIO"))
        query_saidas_anterior_custeio_ = query_saidas.filter(saidas.ano <= ano_anterior, saidas.capcus.like("CUSTEIO"))
        query_entradas_anterior_capital_ = query_entradas.filter(entradas.ano <= ano_anterior, entradas.capcus.like("CAPITAL")) 
        query_saidas_anterior_capital_ = query_saidas.filter(saidas.ano <= ano_anterior, saidas.capcus.like("CAPITAL"))      

        entradas_lista_anterior_custeio = query_entradas_anterior_custeio.all()
        entradas_lista_anterior_capital = query_entradas_anterior_capital.all()
        saidas_lista_anterior_custeio = query_saidas_anterior_custeio.all()
        saidas_lista_anterior_capital = query_saidas_anterior_capital.all()

        entradas_lista_anterior_custeio_ = query_entradas_anterior_custeio_.all()
        entradas_lista_anterior_capital_ = query_entradas_anterior_capital_.all()
        saidas_lista_anterior_custeio_ = query_saidas_anterior_custeio_.all()
        saidas_lista_anterior_capital_ = query_saidas_anterior_capital_.all()
    
        total_entradas_anterior_custeio = sum([entrada.valor for entrada in entradas_lista_anterior_custeio])
        total_entradas_anterior_capital = sum([entrada.valor for entrada in entradas_lista_anterior_capital])
        total_saidas_anterior_custeio = sum([saida.valor1 for saida in saidas_lista_anterior_custeio])
        total_saidas_anterior_capital = sum([saida.valor1 for saida in saidas_lista_anterior_capital])

        total_entradas_anterior_custeio_ = sum([entrada.valor for entrada in entradas_lista_anterior_custeio_])
        total_entradas_anterior_capital_ = sum([entrada.valor for entrada in entradas_lista_anterior_capital_])
        total_saidas_anterior_custeio_ = sum([saida.valor1 for saida in saidas_lista_anterior_custeio_])
        total_saidas_anterior_capital_ = sum([saida.valor1 for saida in saidas_lista_anterior_capital_])

        saldo_anterior_custeio = total_entradas_anterior_custeio + total_entradas_anterior_custeio_ - total_saidas_anterior_custeio - total_saidas_anterior_custeio_
        saldo_anterior_capital = total_entradas_anterior_capital +total_entradas_anterior_capital_ - total_saidas_anterior_capital - total_saidas_anterior_capital_

        saldo_anterior_custeio = saldo_anterior_custeio/100
        saldo_anterior_capital = saldo_anterior_capital/100
        
        #Recurso financeiro recebido
        query_entradas_recebido_custeio = query_entradas.filter(entradas.ano == ano_int, entradas.mes > 6,  entradas.capcus.like("CUSTEIO"))
        query_entradas_recebido_capital = query_entradas.filter(entradas.ano == ano_int, entradas.mes > 6, entradas.capcus.like("CAPITAL"))
        
        entradas_lista_recebido_custeio = query_entradas_recebido_custeio.all()
        entradas_lista_recebido_capital = query_entradas_recebido_capital.all()

        total_entradas_recebido_custeio = sum([entrada.valor for entrada in entradas_lista_recebido_custeio])
        total_entradas_recebido_capital = sum([entrada.valor for entrada in entradas_lista_recebido_capital])

        total_entradas_recebido_custeio = total_entradas_recebido_custeio/100
        total_entradas_recebido_capital = total_entradas_recebido_capital/100
        
        #Recursos próprios
        query_entradas_proprio_custeio = query_entradas.filter(entradas.ano == ano_int, entradas.mes > 6, entradas.comentario == "RECURSOS PRÓPRIOS", entradas.capcus.like("CUSTEIO"))
        query_entradas_proprio_capital = query_entradas.filter(entradas.ano == ano_int, entradas.mes > 6, entradas.comentario == "RECURSOS PRÓPRIOS", entradas.capcus.like("CAPITAL"))
        
        entradas_lista_proprio_custeio = query_entradas_proprio_custeio.all()
        entradas_lista_proprio_capital = query_entradas_proprio_capital.all()

        total_entradas_proprio_custeio = sum([entrada.valor for entrada in entradas_lista_proprio_custeio])
        total_entradas_proprio_capital = sum([entrada.valor for entrada in entradas_lista_proprio_capital])

        total_entradas_proprio_custeio = total_entradas_proprio_custeio/100
        total_entradas_proprio_capital = total_entradas_proprio_capital/100

        #Rendimentos
        query_entradas_rendimento_custeio = query_entradas.filter(entradas.ano == ano_int, entradas.mes > 6, entradas.comentario == "RENDIMENTOS", entradas.capcus.like("CUSTEIO"))
        query_entradas_rendimento_capital = query_entradas.filter(entradas.ano == ano_int, entradas.mes > 6, entradas.comentario == "RENDIMENTOS", entradas.capcus.like("CAPITAL"))
        
        entradas_lista_rendimento_custeio = query_entradas_rendimento_custeio.all()
        entradas_lista_rendimento_capital = query_entradas_rendimento_capital.all()

        total_entradas_rendimento_custeio = sum([entrada.valor for entrada in entradas_lista_rendimento_custeio])
        total_entradas_rendimento_capital = sum([entrada.valor for entrada in entradas_lista_rendimento_capital])

        total_entradas_rendimento_custeio = total_entradas_rendimento_custeio/100
        total_entradas_rendimento_capital = total_entradas_rendimento_capital/100

        #Recurso calculado


        total_entradas_recebido_custeio = total_entradas_recebido_custeio - total_entradas_proprio_custeio - total_entradas_rendimento_custeio
        total_entradas_recebido_capital = total_entradas_recebido_capital - total_entradas_proprio_capital - total_entradas_rendimento_capital

        #Receita total

        total_receita_custeio = saldo_anterior_custeio + total_entradas_recebido_custeio + total_entradas_proprio_custeio + total_entradas_rendimento_custeio

        total_receita_capital = saldo_anterior_capital + total_entradas_recebido_capital + total_entradas_proprio_capital + total_entradas_rendimento_capital

        #Despesas

        query_saidas_custeio = query_saidas.filter(saidas.ano == ano_int, saidas.mes > 6, saidas.capcus.like("CUSTEIO"))
        query_saidas_capital = query_saidas.filter(saidas.ano == ano_int, saidas.mes > 6, saidas.capcus.like("CAPITAL"))

        saidas_lista_custeio = query_saidas_custeio.all()
        saidas_lista_capital = query_saidas_capital.all()

        total_saidas_custeio = sum([saida.valor1 for saida in saidas_lista_custeio])
        total_saidas_capital = sum([saida.valor1 for saida in saidas_lista_capital])

        total_saidas_custeio = total_saidas_custeio/100
        total_saidas_capital = total_saidas_capital/100

        #Saldo final

        saldo_final_custeio = total_receita_custeio - total_saidas_custeio
        saldo_final_capital = total_receita_capital - total_saidas_capital
            
    

    # Criar o documento Word
    doc = Document()
    temp_dir = tempfile.gettempdir()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1)       # Margem superior de 1 cm
        section.bottom_margin = Cm(1)    # Margem inferior de 1 cm
        section.left_margin = Cm(1)      # Margem esquerda de 1 cm
        section.right_margin = Cm(1)     # Margem direita de 1 cm
    # Configurar estilo de espaçamento simples globalmente
    style = doc.styles['Normal']
    style.font.size = Pt(10)  # Define tamanho padrão da fonte
    style.paragraph_format.line_spacing = 1.0  # Define espaçamento simples

    dados_tabela = [
        ['UNIDADE EXECUTORA:', 'CNPJ:'],
        [f'CONSELHO ESCOLAR DA {escola_data.nome_escola.upper()}', escola_data.cnpj_conselho],
        ['RECURSO:', 'EXERCÍCIO:'],
        [fonte_dados, dadoperiodo]
    ]
    
    tabela = doc.add_table(rows=len(dados_tabela), cols=len(dados_tabela[0]))

    for i, linha in enumerate(dados_tabela):
        for j, celula in enumerate(linha):
            tabela.cell(i, j).text = str(celula)
    
    tabela.style = 'Table Grid'

    doc.add_heading("PARECER DO CONSELHO FISCAL", level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Adicionar informações da reunião
    
    doc.add_paragraph(f"O Conselho Fiscal desta Unidade Executora recomenda a aprovação da prestação de contas dos recursos do Programa {fonte_dados}, referente ao período compreendido entre 01/{mesinicial}/{ano_dados} e {diafinal}/{mesfinal}/{ano_dados}, conforme demonstração abaixo:").alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    

    dados_tabela_valores = [
        ['', 'CUSTEIO', 'CAPITAL'],
        ['Saldo do Exercício Anterior', formatar_valor(saldo_anterior_custeio), formatar_valor(saldo_anterior_capital)],
        ['Recurso Financeiro Recebido', formatar_valor(total_entradas_recebido_custeio), formatar_valor(total_entradas_recebido_capital)],
        ['Recursos Próprios', formatar_valor(total_entradas_proprio_custeio), formatar_valor(total_entradas_proprio_capital)],
        ['Rendimentos Provenientes de Aplicações Financeiras', formatar_valor(total_entradas_rendimento_custeio), formatar_valor(total_entradas_rendimento_capital)],
        ['(-) Recursos Devolvidos à Conta da Mantenedora', 'R$ 0,00', 'R$ 0,00'],
        ['Receita Total', formatar_valor(total_receita_custeio), formatar_valor(total_receita_capital)],
        ['Despesa Realizada', formatar_valor(total_saidas_custeio), formatar_valor(total_saidas_capital)],
        ['Saldo Final', formatar_valor(saldo_final_custeio), formatar_valor(saldo_final_capital)]
    ]
    
    tabela_valores = doc.add_table(rows=len(dados_tabela_valores), cols=len(dados_tabela_valores[0]))

    for i, linha in enumerate(dados_tabela_valores):
        for j, celula in enumerate(linha):
            tabela_valores.cell(i, j).text = str(celula)
    
    tabela_valores.style = 'Table Grid'
    
    paragrafo_espaco = doc.add_paragraph()
    paragrafo_espaco.paragraph_format.space_after = Pt(20)
    
    
    # Assinaturas
    doc.add_paragraph("\n1º Conselheiro Fiscal              _____________________________________________________________________________________").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f"NOME: {escola_data.nome_1_conselheiro}").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f"ENDEREÇO: {escola_data.endereco_1_conselheiro}").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f"CPF: {escola_data.cpf_1_conselheiro}").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragrafo_espaco.paragraph_format.space_after = Pt(20)
    doc.add_paragraph("\n2º Conselheiro Fiscal              _____________________________________________________________________________________").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f"NOME: {escola_data.nome_2_conselheiro}").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f"ENDEREÇO: {escola_data.endereco_2_conselheiro}").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f"CPF: {escola_data.cpf_2_conselheiro}").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragrafo_espaco.paragraph_format.space_after = Pt(20)
    doc.add_paragraph("\n3º Conselheiro Fiscal              _____________________________________________________________________________________").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f"NOME: {escola_data.nome_3_conselheiro}").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f"ENDEREÇO: {escola_data.endereco_3_conselheiro}").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f"CPF: {escola_data.cpf_3_conselheiro}").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragrafo_espaco.paragraph_format.space_after = Pt(20)
    doc.add_paragraph("\n4º Conselheiro Fiscal              _____________________________________________________________________________________").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f"NOME: {escola_data.nome_4_conselheiro}").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f"ENDEREÇO: {escola_data.endereco_4_conselheiro}").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f"CPF: {escola_data.cpf_4_conselheiro}").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragrafo_espaco.paragraph_format.space_after = Pt(20)
    
       

    # Salvar o arquivo em um diretório temporário
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, f"parecer_{fonte_dados}_{ano_dados}.docx")
    doc.save(file_path)

    # Retornar o arquivo Word gerado
    return send_file(file_path, as_attachment=True, download_name=f"parecer_{fonte_dados}_{ano_dados}.docx")
   

@app.route('/gerar_oficio', methods=['GET'])
def gerar_oficio():
    ano_dados = request.args.get('ano')
    fonte_dados=request.args.get('fonte')
    mandato_dados = request.args.get('anoMandato')
    data_dados = request.args.get('data')
    mes_dados = request.args.get('mes')
    escola_data = Escola.query.filter_by(id=1).first()

    escola_data = Escola.query.filter_by(ano_mandato=mandato_dados).first()

    from datetime import datetime
    datadocumento = data_dados
    data_formatada = datetime.strptime(datadocumento, '%Y-%m-%d').strftime('%d de %B de %Y')


    query_entradas = entradas.query
    query_saidas = saidas.query

    query_entradas = query_entradas.filter_by(fonte=fonte_dados)
    query_saidas = query_saidas.filter_by(fonte=fonte_dados)

    
    ano_int = int(ano_dados)
    ano_anterior = ano_int-1
    periodo = int(mes_dados)
    if periodo == 3:
        dadoperiodo = " "
        #Saldo do exercício anterior
        query_entradas_anterior_custeio = query_entradas.filter(entradas.ano <= ano_anterior, entradas.capcus.like("CUSTEIO"))
        query_saidas_anterior_custeio = query_saidas.filter(saidas.ano <= ano_anterior, saidas.capcus.like("CUSTEIO"))
        query_entradas_anterior_capital = query_entradas.filter(entradas.ano <= ano_anterior, entradas.capcus.like("CAPITAL")) 
        query_saidas_anterior_capital = query_saidas.filter(saidas.ano <= ano_anterior, saidas.capcus.like("CAPITAL"))    

        entradas_lista_anterior_custeio = query_entradas_anterior_custeio.all()
        entradas_lista_anterior_capital = query_entradas_anterior_capital.all()
        saidas_lista_anterior_custeio = query_saidas_anterior_custeio.all()
        saidas_lista_anterior_capital = query_saidas_anterior_capital.all()
    
        total_entradas_anterior_custeio = sum([entrada.valor for entrada in entradas_lista_anterior_custeio])
        total_entradas_anterior_capital = sum([entrada.valor for entrada in entradas_lista_anterior_capital])
        total_saidas_anterior_custeio = sum([saida.valor1 for saida in saidas_lista_anterior_custeio])
        total_saidas_anterior_capital = sum([saida.valor1 for saida in saidas_lista_anterior_capital])

        saldo_anterior_custeio = total_entradas_anterior_custeio - total_saidas_anterior_custeio
        saldo_anterior_capital = total_entradas_anterior_capital - total_saidas_anterior_capital

        saldo_anterior_custeio = saldo_anterior_custeio/100
        saldo_anterior_capital = saldo_anterior_capital/100
    
        #Recurso financeiro recebido
        query_entradas_recebido_custeio = query_entradas.filter(entradas.ano == ano_int,  entradas.capcus.like("CUSTEIO"))
        query_entradas_recebido_capital = query_entradas.filter(entradas.ano == ano_int, entradas.capcus.like("CAPITAL"))
    
        entradas_lista_recebido_custeio = query_entradas_recebido_custeio.all()
        entradas_lista_recebido_capital = query_entradas_recebido_capital.all()

        total_entradas_recebido_custeio = sum([entrada.valor for entrada in entradas_lista_recebido_custeio])
        total_entradas_recebido_capital = sum([entrada.valor for entrada in entradas_lista_recebido_capital])

        total_entradas_recebido_custeio = total_entradas_recebido_custeio/100
        total_entradas_recebido_capital = total_entradas_recebido_capital/100
    
        #Recursos próprios
        query_entradas_proprio_custeio = query_entradas.filter(entradas.ano == ano_int, entradas.comentario == "RECURSOS PRÓPRIOS", entradas.capcus.like("CUSTEIO"))
        query_entradas_proprio_capital = query_entradas.filter(entradas.ano == ano_int, entradas.comentario == "RECURSOS PRÓPRIOS", entradas.capcus.like("CAPITAL"))
    
        entradas_lista_proprio_custeio = query_entradas_proprio_custeio.all()
        entradas_lista_proprio_capital = query_entradas_proprio_capital.all()

        total_entradas_proprio_custeio = sum([entrada.valor for entrada in entradas_lista_proprio_custeio])
        total_entradas_proprio_capital = sum([entrada.valor for entrada in entradas_lista_proprio_capital])

        total_entradas_proprio_custeio = total_entradas_proprio_custeio/100
        total_entradas_proprio_capital = total_entradas_proprio_capital/100

        #Rendimentos
        query_entradas_rendimento_custeio = query_entradas.filter(entradas.ano == ano_int, entradas.comentario == "RENDIMENTOS", entradas.capcus.like("CUSTEIO"))
        query_entradas_rendimento_capital = query_entradas.filter(entradas.ano == ano_int, entradas.comentario == "RENDIMENTOS", entradas.capcus.like("CAPITAL"))
    
        entradas_lista_rendimento_custeio = query_entradas_rendimento_custeio.all()
        entradas_lista_rendimento_capital = query_entradas_rendimento_capital.all()

        total_entradas_rendimento_custeio = sum([entrada.valor for entrada in entradas_lista_rendimento_custeio])
        total_entradas_rendimento_capital = sum([entrada.valor for entrada in entradas_lista_rendimento_capital])

        total_entradas_rendimento_custeio = total_entradas_rendimento_custeio/100
        total_entradas_rendimento_capital = total_entradas_rendimento_capital/100

        #Recurso calculado


        total_entradas_recebido_custeio = total_entradas_recebido_custeio - total_entradas_proprio_custeio - total_entradas_rendimento_custeio
        total_entradas_recebido_capital = total_entradas_recebido_capital - total_entradas_proprio_capital - total_entradas_rendimento_capital

        #Receita total

        total_receita_custeio = saldo_anterior_custeio + total_entradas_recebido_custeio + total_entradas_proprio_custeio + total_entradas_rendimento_custeio

        total_receita_capital = saldo_anterior_capital + total_entradas_recebido_capital + total_entradas_proprio_capital + total_entradas_rendimento_capital

        #Despesas

        query_saidas_custeio = query_saidas.filter(saidas.ano == ano_int, saidas.capcus.like("CUSTEIO"))
        query_saidas_capital = query_saidas.filter(saidas.ano == ano_int, saidas.capcus.like("CAPITAL"))

        saidas_lista_custeio = query_saidas_custeio.all()
        saidas_lista_capital = query_saidas_capital.all()

        total_saidas_custeio = sum([saida.valor1 for saida in saidas_lista_custeio])
        total_saidas_capital = sum([saida.valor1 for saida in saidas_lista_capital])

        total_saidas_custeio = total_saidas_custeio/100
        total_saidas_capital = total_saidas_capital/100

        #Saldo final

        saldo_final_custeio = total_receita_custeio - total_saidas_custeio
        saldo_final_capital = total_receita_capital - total_saidas_capital

        receitatotal = total_receita_custeio + total_receita_capital

        valor_extenso = valor_por_extenso(receitatotal)

    if periodo == 2:
        dadoperiodo = ".2"
        query_entradas_anterior_custeio = query_entradas.filter(entradas.ano == ano_int, entradas.mes <= 6, entradas.capcus.like("CUSTEIO"))
        query_saidas_anterior_custeio = query_saidas.filter(saidas.ano == ano_int, saidas.mes <= 6, saidas.capcus.like("CUSTEIO"))
        query_entradas_anterior_capital = query_entradas.filter(entradas.ano == ano_int, entradas.mes <= 6, entradas.capcus.like("CAPITAL")) 
        query_saidas_anterior_capital = query_saidas.filter(saidas.ano == ano_int, saidas.mes <= 6, saidas.capcus.like("CAPITAL"))

        query_entradas_anterior_custeio_ = query_entradas.filter(entradas.ano <= ano_anterior, entradas.capcus.like("CUSTEIO"))
        query_saidas_anterior_custeio_ = query_saidas.filter(saidas.ano <= ano_anterior, saidas.capcus.like("CUSTEIO"))
        query_entradas_anterior_capital_ = query_entradas.filter(entradas.ano <= ano_anterior, entradas.capcus.like("CAPITAL")) 
        query_saidas_anterior_capital_ = query_saidas.filter(saidas.ano <= ano_anterior, saidas.capcus.like("CAPITAL"))      

        entradas_lista_anterior_custeio = query_entradas_anterior_custeio.all()
        entradas_lista_anterior_capital = query_entradas_anterior_capital.all()
        saidas_lista_anterior_custeio = query_saidas_anterior_custeio.all()
        saidas_lista_anterior_capital = query_saidas_anterior_capital.all()

        entradas_lista_anterior_custeio_ = query_entradas_anterior_custeio_.all()
        entradas_lista_anterior_capital_ = query_entradas_anterior_capital_.all()
        saidas_lista_anterior_custeio_ = query_saidas_anterior_custeio_.all()
        saidas_lista_anterior_capital_ = query_saidas_anterior_capital_.all()
    
        total_entradas_anterior_custeio = sum([entrada.valor for entrada in entradas_lista_anterior_custeio])
        total_entradas_anterior_capital = sum([entrada.valor for entrada in entradas_lista_anterior_capital])
        total_saidas_anterior_custeio = sum([saida.valor1 for saida in saidas_lista_anterior_custeio])
        total_saidas_anterior_capital = sum([saida.valor1 for saida in saidas_lista_anterior_capital])

        total_entradas_anterior_custeio_ = sum([entrada.valor for entrada in entradas_lista_anterior_custeio_])
        total_entradas_anterior_capital_ = sum([entrada.valor for entrada in entradas_lista_anterior_capital_])
        total_saidas_anterior_custeio_ = sum([saida.valor1 for saida in saidas_lista_anterior_custeio_])
        total_saidas_anterior_capital_ = sum([saida.valor1 for saida in saidas_lista_anterior_capital_])

        saldo_anterior_custeio = total_entradas_anterior_custeio + total_entradas_anterior_custeio_ - total_saidas_anterior_custeio - total_saidas_anterior_custeio_
        saldo_anterior_capital = total_entradas_anterior_capital +total_entradas_anterior_capital_ - total_saidas_anterior_capital - total_saidas_anterior_capital_

        saldo_anterior_custeio = saldo_anterior_custeio/100
        saldo_anterior_capital = saldo_anterior_capital/100
    
        #Recurso financeiro recebido
        query_entradas_recebido_custeio = query_entradas.filter(entradas.ano == ano_int, entradas.mes > 6,  entradas.capcus.like("CUSTEIO"))
        query_entradas_recebido_capital = query_entradas.filter(entradas.ano == ano_int, entradas.mes > 6, entradas.capcus.like("CAPITAL"))
    
        entradas_lista_recebido_custeio = query_entradas_recebido_custeio.all()
        entradas_lista_recebido_capital = query_entradas_recebido_capital.all()

        total_entradas_recebido_custeio = sum([entrada.valor for entrada in entradas_lista_recebido_custeio])
        total_entradas_recebido_capital = sum([entrada.valor for entrada in entradas_lista_recebido_capital])

        total_entradas_recebido_custeio = total_entradas_recebido_custeio/100
        total_entradas_recebido_capital = total_entradas_recebido_capital/100
    
        #Recursos próprios
        query_entradas_proprio_custeio = query_entradas.filter(entradas.ano == ano_int, entradas.mes > 6, entradas.comentario == "RECURSOS PRÓPRIOS", entradas.capcus.like("CUSTEIO"))
        query_entradas_proprio_capital = query_entradas.filter(entradas.ano == ano_int, entradas.mes > 6, entradas.comentario == "RECURSOS PRÓPRIOS", entradas.capcus.like("CAPITAL"))
    
        entradas_lista_proprio_custeio = query_entradas_proprio_custeio.all()
        entradas_lista_proprio_capital = query_entradas_proprio_capital.all()

        total_entradas_proprio_custeio = sum([entrada.valor for entrada in entradas_lista_proprio_custeio])
        total_entradas_proprio_capital = sum([entrada.valor for entrada in entradas_lista_proprio_capital])

        total_entradas_proprio_custeio = total_entradas_proprio_custeio/100
        total_entradas_proprio_capital = total_entradas_proprio_capital/100

        #Rendimentos
        query_entradas_rendimento_custeio = query_entradas.filter(entradas.ano == ano_int, entradas.mes > 6, entradas.comentario == "RENDIMENTOS", entradas.capcus.like("CUSTEIO"))
        query_entradas_rendimento_capital = query_entradas.filter(entradas.ano == ano_int, entradas.mes > 6, entradas.comentario == "RENDIMENTOS", entradas.capcus.like("CAPITAL"))
    
        entradas_lista_rendimento_custeio = query_entradas_rendimento_custeio.all()
        entradas_lista_rendimento_capital = query_entradas_rendimento_capital.all()

        total_entradas_rendimento_custeio = sum([entrada.valor for entrada in entradas_lista_rendimento_custeio])
        total_entradas_rendimento_capital = sum([entrada.valor for entrada in entradas_lista_rendimento_capital])

        total_entradas_rendimento_custeio = total_entradas_rendimento_custeio/100
        total_entradas_rendimento_capital = total_entradas_rendimento_capital/100

        #Recurso calculado


        total_entradas_recebido_custeio = total_entradas_recebido_custeio - total_entradas_proprio_custeio - total_entradas_rendimento_custeio
        total_entradas_recebido_capital = total_entradas_recebido_capital - total_entradas_proprio_capital - total_entradas_rendimento_capital

        #Receita total

        total_receita_custeio = saldo_anterior_custeio + total_entradas_recebido_custeio + total_entradas_proprio_custeio + total_entradas_rendimento_custeio

        total_receita_capital = saldo_anterior_capital + total_entradas_recebido_capital + total_entradas_proprio_capital + total_entradas_rendimento_capital

        #Despesas

        query_saidas_custeio = query_saidas.filter(saidas.ano == ano_int, saidas.mes > 6, saidas.capcus.like("CUSTEIO"))
        query_saidas_capital = query_saidas.filter(saidas.ano == ano_int, saidas.mes > 6, saidas.capcus.like("CAPITAL"))

        saidas_lista_custeio = query_saidas_custeio.all()
        saidas_lista_capital = query_saidas_capital.all()

        total_saidas_custeio = sum([saida.valor1 for saida in saidas_lista_custeio])
        total_saidas_capital = sum([saida.valor1 for saida in saidas_lista_capital])

        total_saidas_custeio = total_saidas_custeio/100
        total_saidas_capital = total_saidas_capital/100

        #Saldo final

        saldo_final_custeio = total_receita_custeio - total_saidas_custeio
        saldo_final_capital = total_receita_capital - total_saidas_capital

        receitatotal = total_receita_custeio + total_receita_capital

        valor_extenso = valor_por_extenso(receitatotal)
    
    if periodo == 1:
        dadoperiodo = ".1"
        #Saldo do exercício anterior
        query_entradas_anterior_custeio = query_entradas.filter(entradas.ano <= ano_anterior, entradas.capcus.like("CUSTEIO"))
        query_saidas_anterior_custeio = query_saidas.filter(saidas.ano <= ano_anterior, saidas.capcus.like("CUSTEIO"))
        query_entradas_anterior_capital = query_entradas.filter(entradas.ano <= ano_anterior, entradas.capcus.like("CAPITAL")) 
        query_saidas_anterior_capital = query_saidas.filter(saidas.ano <= ano_anterior, saidas.capcus.like("CAPITAL"))     

        entradas_lista_anterior_custeio = query_entradas_anterior_custeio.all()
        entradas_lista_anterior_capital = query_entradas_anterior_capital.all()
        saidas_lista_anterior_custeio = query_saidas_anterior_custeio.all()
        saidas_lista_anterior_capital = query_saidas_anterior_capital.all()
    
        total_entradas_anterior_custeio = sum([entrada.valor for entrada in entradas_lista_anterior_custeio])
        total_entradas_anterior_capital = sum([entrada.valor for entrada in entradas_lista_anterior_capital])
        total_saidas_anterior_custeio = sum([saida.valor1 for saida in saidas_lista_anterior_custeio])
        total_saidas_anterior_capital = sum([saida.valor1 for saida in saidas_lista_anterior_capital])

        saldo_anterior_custeio = total_entradas_anterior_custeio - total_saidas_anterior_custeio
        saldo_anterior_capital = total_entradas_anterior_capital - total_saidas_anterior_capital

        saldo_anterior_custeio = saldo_anterior_custeio/100
        saldo_anterior_capital = saldo_anterior_capital/100
    
        #Recurso financeiro recebido
        query_entradas_recebido_custeio = query_entradas.filter(entradas.ano == ano_int, entradas.mes <= 6,  entradas.capcus.like("CUSTEIO"))
        query_entradas_recebido_capital = query_entradas.filter(entradas.ano == ano_int, entradas.mes <= 6, entradas.capcus.like("CAPITAL"))
    
        entradas_lista_recebido_custeio = query_entradas_recebido_custeio.all()
        entradas_lista_recebido_capital = query_entradas_recebido_capital.all()

        total_entradas_recebido_custeio = sum([entrada.valor for entrada in entradas_lista_recebido_custeio])
        total_entradas_recebido_capital = sum([entrada.valor for entrada in entradas_lista_recebido_capital])

        total_entradas_recebido_custeio = total_entradas_recebido_custeio/100
        total_entradas_recebido_capital = total_entradas_recebido_capital/100
    
        #Recursos próprios
        query_entradas_proprio_custeio = query_entradas.filter(entradas.ano == ano_int, entradas.mes <= 6, entradas.comentario == "RECURSOS PRÓPRIOS", entradas.capcus.like("CUSTEIO"))
        query_entradas_proprio_capital = query_entradas.filter(entradas.ano == ano_int, entradas.mes <= 6, entradas.comentario == "RECURSOS PRÓPRIOS", entradas.capcus.like("CAPITAL"))
    
        entradas_lista_proprio_custeio = query_entradas_proprio_custeio.all()
        entradas_lista_proprio_capital = query_entradas_proprio_capital.all()

        total_entradas_proprio_custeio = sum([entrada.valor for entrada in entradas_lista_proprio_custeio])
        total_entradas_proprio_capital = sum([entrada.valor for entrada in entradas_lista_proprio_capital])

        total_entradas_proprio_custeio = total_entradas_proprio_custeio/100
        total_entradas_proprio_capital = total_entradas_proprio_capital/100

        #Rendimentos
        query_entradas_rendimento_custeio = query_entradas.filter(entradas.ano == ano_int, entradas.mes <= 6, entradas.comentario == "RENDIMENTOS", entradas.capcus.like("CUSTEIO"))
        query_entradas_rendimento_capital = query_entradas.filter(entradas.ano == ano_int, entradas.mes <= 6, entradas.comentario == "RENDIMENTOS", entradas.capcus.like("CAPITAL"))
    
        entradas_lista_rendimento_custeio = query_entradas_rendimento_custeio.all()
        entradas_lista_rendimento_capital = query_entradas_rendimento_capital.all()

        total_entradas_rendimento_custeio = sum([entrada.valor for entrada in entradas_lista_rendimento_custeio])
        total_entradas_rendimento_capital = sum([entrada.valor for entrada in entradas_lista_rendimento_capital])

        total_entradas_rendimento_custeio = total_entradas_rendimento_custeio/100
        total_entradas_rendimento_capital = total_entradas_rendimento_capital/100

        #Recurso calculado


        total_entradas_recebido_custeio = total_entradas_recebido_custeio - total_entradas_proprio_custeio - total_entradas_rendimento_custeio
        total_entradas_recebido_capital = total_entradas_recebido_capital - total_entradas_proprio_capital - total_entradas_rendimento_capital

        #Receita total

        total_receita_custeio = saldo_anterior_custeio + total_entradas_recebido_custeio + total_entradas_proprio_custeio + total_entradas_rendimento_custeio

        total_receita_capital = saldo_anterior_capital + total_entradas_recebido_capital + total_entradas_proprio_capital + total_entradas_rendimento_capital

        #Despesas

        query_saidas_custeio = query_saidas.filter(saidas.ano == ano_int, saidas.mes <= 6, saidas.capcus.like("CUSTEIO"))
        query_saidas_capital = query_saidas.filter(saidas.ano == ano_int, saidas.mes <= 6, saidas.capcus.like("CAPITAL"))

        saidas_lista_custeio = query_saidas_custeio.all()
        saidas_lista_capital = query_saidas_capital.all()

        total_saidas_custeio = sum([saida.valor1 for saida in saidas_lista_custeio])
        total_saidas_capital = sum([saida.valor1 for saida in saidas_lista_capital])

        total_saidas_custeio = total_saidas_custeio/100
        total_saidas_capital = total_saidas_capital/100

        #Saldo final

        saldo_final_custeio = total_receita_custeio - total_saidas_custeio
        saldo_final_capital = total_receita_capital - total_saidas_capital

        receitatotal = total_receita_custeio + total_receita_capital

        valor_extenso = valor_por_extenso(receitatotal)        
    


    # Criar o documento Word
    doc = Document()
    temp_dir = tempfile.gettempdir()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1)       # Margem superior de 1 cm
        section.bottom_margin = Cm(1)    # Margem inferior de 1 cm
        section.left_margin = Cm(1)      # Margem esquerda de 1 cm
        section.right_margin = Cm(1)     # Margem direita de 1 cm
    # Configurar estilo de espaçamento simples globalmente
    style = doc.styles['Normal']
    style.font.size = Pt(10)  # Define tamanho padrão da fonte
    style.paragraph_format.line_spacing = 1.0  # Define espaçamento simples

    

    doc.add_heading(f"UNIDADE EXECUTORA: CONSELHO ESCOLAR DA {escola_data.nome_escola.upper()}"  , level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Adicionar informações da reunião
    
    doc.add_paragraph("OF. NÚMERO _____________").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph(f"{escola_data.cidade} - AL, {data_formatada}. ").alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.add_paragraph("Senhor(a) Secretário(a)").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph(f"Encaminha-se a prestação de contas do Programa {fonte_dados} - {ano_dados}.{dadoperiodo}, cuja programação de despesas teve aprovação de Vossa Excelência.").alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_paragraph(f"Vale ressaltar que a receita total, correspondente à R$ {formatar_valor(receitatotal)} ({valor_extenso}), destinada a custear despesas gerais para manutenção desta Unidade Executora, foi obtida conforme demonstração abaixo:").alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    

    dados_tabela_valores = [
        ['DESCRIÇÃO DA RECEITA', 'CUSTEIO', 'CAPITAL'],
        ['Saldo do Exercício Anterior', formatar_valor(saldo_anterior_custeio), formatar_valor(saldo_anterior_capital)],
        ['Recurso Financeiro da Mantenedoura', formatar_valor(total_entradas_recebido_custeio), formatar_valor(total_entradas_recebido_capital)],
        ['Recursos Próprios', formatar_valor(total_entradas_proprio_custeio), formatar_valor(total_entradas_proprio_capital)],
        ['Rendimentos Provenientes de Aplicações Financeiras', formatar_valor(total_entradas_rendimento_custeio), formatar_valor(total_entradas_rendimento_capital)],
        ['(-) Recursos Devolvidos à Conta da Mantenedora', 'R$ 0,00', 'R$ 0,00'],
        ['Receita Total', formatar_valor(total_receita_custeio), formatar_valor(total_receita_capital)],
        ['Receita Total Consolidada', formatar_valor(total_receita_custeio), formatar_valor(total_receita_capital)]
    ]
    
    tabela_valores = doc.add_table(rows=len(dados_tabela_valores), cols=len(dados_tabela_valores[0]))

    for i, linha in enumerate(dados_tabela_valores):
        for j, celula in enumerate(linha):
            tabela_valores.cell(i, j).text = str(celula)
    
    tabela_valores.style = 'Table Grid'
    
    paragrafo_espaco = doc.add_paragraph()
    paragrafo_espaco.paragraph_format.space_after = Pt(20)
    
    
    # Assinaturas
    doc.add_paragraph("Na oportunidade reiteramos votos de elevada estima e consideração.").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    paragrafo_espaco.paragraph_format.space_after = Pt(30)

    doc.add_paragraph("Atenciosamente").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragrafo_espaco.paragraph_format.space_after = Pt(30)

    doc.add_paragraph("________________________________________________").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("ASSINATURA DO REPRESENTANTE DA UNIDADE EXECUTORA").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragrafo_espaco.paragraph_format.space_after = Pt(30)

    doc.add_paragraph(f"{escola_data.presidente_conselho}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(f"NOME DO(A) PRESIDENTE DA UNIDADE EXECUTORA").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
       

    # Salvar o arquivo em um diretório temporário
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, f"parecer_{fonte_dados}_{ano_dados}.docx")
    doc.save(file_path)

    # Retornar o arquivo Word gerado
    return send_file(file_path, as_attachment=True, download_name=f"parecer_{fonte_dados}_{ano_dados}.docx")
   

@app.route('/gerar_demonstrativo', methods=['GET'])
def gerar_demonstrativo():
    # Obtenção dos parâmetros da requisição
    ano_dados = request.args.get('ano')
    fonte_dados = request.args.get('fonte')
    mandato_dados = request.args.get('anoMandato')
    data_dados = request.args.get('data') # Não utilizado no código, mas mantido
    mes_dados = request.args.get('mes')

    # Busca de dados da escola
    escola_data = Escola.query.filter_by(ano_mandato=mandato_dados).first()

    # Queries base para entradas e saídas
    query_entradas = entradas.query.filter_by(fonte=fonte_dados)
    query_saidas = saidas.query.filter_by(fonte=fonte_dados)

    ano_int = int(ano_dados)
    ano_anterior = ano_int - 1
    periodo = int(mes_dados)

    # --- Lógica para o ANO INTEIRO (periodo == 3) ---
    if periodo == 3:
        # Saldo do exercício anterior
        query_entradas_anterior_custeio = query_entradas.filter(entradas.ano <= ano_anterior, entradas.capcus.like("CUSTEIO"))
        query_saidas_anterior_custeio = query_saidas.filter(saidas.ano <= ano_anterior, saidas.capcus.like("CUSTEIO"))
        query_entradas_anterior_capital = query_entradas.filter(entradas.ano <= ano_anterior, entradas.capcus.like("CAPITAL"))
        query_saidas_anterior_capital = query_saidas.filter(saidas.ano <= ano_anterior, saidas.capcus.like("CAPITAL"))

        saldo_anterior_custeio = (sum(e.valor for e in query_entradas_anterior_custeio.all()) - sum(s.valor1 for s in query_saidas_anterior_custeio.all())) / 100
        saldo_anterior_capital = (sum(e.valor for e in query_entradas_anterior_capital.all()) - sum(s.valor1 for s in query_saidas_anterior_capital.all())) / 100

        # Entradas do ano atual
        entradas_ano_custeio = query_entradas.filter(entradas.ano == ano_int, entradas.capcus.like("CUSTEIO")).all()
        entradas_ano_capital = query_entradas.filter(entradas.ano == ano_int, entradas.capcus.like("CAPITAL")).all()

        total_entradas_proprio_custeio = sum(e.valor for e in entradas_ano_custeio if e.comentario == "RECURSOS PRÓPRIOS") / 100
        total_entradas_proprio_capital = sum(e.valor for e in entradas_ano_capital if e.comentario == "RECURSOS PRÓPRIOS") / 100
        total_entradas_rendimento_custeio = sum(e.valor for e in entradas_ano_custeio if e.comentario == "RENDIMENTOS") / 100
        total_entradas_rendimento_capital = sum(e.valor for e in entradas_ano_capital if e.comentario == "RENDIMENTOS") / 100
        total_entradas_recebido_custeio = (sum(e.valor for e in entradas_ano_custeio) / 100) - total_entradas_proprio_custeio - total_entradas_rendimento_custeio
        total_entradas_recebido_capital = (sum(e.valor for e in entradas_ano_capital) / 100) - total_entradas_proprio_capital - total_entradas_rendimento_capital

        # Receita total
        total_receita_custeio = saldo_anterior_custeio + total_entradas_recebido_custeio + total_entradas_proprio_custeio + total_entradas_rendimento_custeio
        total_receita_capital = saldo_anterior_capital + total_entradas_recebido_capital + total_entradas_proprio_capital + total_entradas_rendimento_capital

        # Despesas
        saidas_lista_custeio = query_saidas.filter(saidas.ano == ano_int, saidas.capcus.like("CUSTEIO")).all()
        saidas_lista_capital = query_saidas.filter(saidas.ano == ano_int, saidas.capcus.like("CAPITAL")).all()
        total_saidas_custeio = sum(s.valor1 for s in saidas_lista_custeio) / 100
        total_saidas_capital = sum(s.valor1 for s in saidas_lista_capital) / 100

        # Saldo final
        saldo_final_custeio = total_receita_custeio - total_saidas_custeio
        saldo_final_capital = total_receita_capital - total_saidas_capital

        # Juntar as listas de saídas
        todas_saidas = saidas_lista_custeio + saidas_lista_capital

    # --- Lógica para o 2º SEMESTRE (periodo == 2) ---
    elif periodo == 2:
        # Lógica de cálculo similar, mas filtrando por mes > 6
        # (O código original tinha uma lógica complexa para saldo anterior, simplificada aqui para clareza)
        # Saldo anterior (até o final do 1º semestre)
        

        entradas_ate_jun_custeio = (query_entradas.filter(entradas.ano <= ano_anterior, entradas.capcus.like("CUSTEIO")).all() + query_entradas.filter(entradas.ano == ano_int, entradas.mes <= 6, entradas.capcus.like("CUSTEIO")).all())
        saidas_ate_jun_custeio = (query_saidas.filter(saidas.ano <= ano_anterior, saidas.capcus.like("CUSTEIO")).all() + query_saidas.filter(saidas.ano == ano_int, saidas.mes <= 6, saidas.capcus.like("CUSTEIO")).all())
        entradas_ate_jun_capital = (query_entradas.filter(entradas.ano <= ano_anterior, entradas.capcus.like("CAPITAL")).all() + query_entradas.filter(entradas.ano == ano_int, entradas.mes <= 6, entradas.capcus.like("CAPITAL")).all())
        saidas_ate_jun_capital = (query_saidas.filter(saidas.ano <= ano_anterior, saidas.capcus.like("CAPITAL")).all() + query_saidas.filter(saidas.ano == ano_int, saidas.mes <= 6, saidas.capcus.like("CAPITAL")).all())

        saldo_anterior_custeio = (sum(e.valor for e in entradas_ate_jun_custeio) - sum(s.valor1 for s in saidas_ate_jun_custeio)) / 100
        saldo_anterior_capital = (sum(e.valor for e in entradas_ate_jun_capital) - sum(s.valor1 for s in saidas_ate_jun_capital)) / 100

        # Entradas do 2º semestre
        entradas_semestre_custeio = query_entradas.filter(entradas.ano == ano_int, entradas.mes > 6, entradas.capcus.like("CUSTEIO")).all()
        entradas_semestre_capital = query_entradas.filter(entradas.ano == ano_int, entradas.mes > 6, entradas.capcus.like("CAPITAL")).all()
        
        total_entradas_proprio_custeio = sum(e.valor for e in entradas_semestre_custeio if e.comentario == "RECURSOS PRÓPRIOS") / 100
        total_entradas_proprio_capital = sum(e.valor for e in entradas_semestre_capital if e.comentario == "RECURSOS PRÓPRIOS") / 100
        total_entradas_rendimento_custeio = sum(e.valor for e in entradas_semestre_custeio if e.comentario == "RENDIMENTOS") / 100
        total_entradas_rendimento_capital = sum(e.valor for e in entradas_semestre_capital if e.comentario == "RENDIMENTOS") / 100
        total_entradas_recebido_custeio = (sum(e.valor for e in entradas_semestre_custeio) / 100) - total_entradas_proprio_custeio - total_entradas_rendimento_custeio
        total_entradas_recebido_capital = (sum(e.valor for e in entradas_semestre_capital) / 100) - total_entradas_proprio_capital - total_entradas_rendimento_capital
        
        total_receita_custeio = saldo_anterior_custeio + total_entradas_recebido_custeio + total_entradas_proprio_custeio + total_entradas_rendimento_custeio
        total_receita_capital = saldo_anterior_capital + total_entradas_recebido_capital + total_entradas_proprio_capital + total_entradas_rendimento_capital

        # Despesas do 2º semestre
        saidas_lista_custeio = query_saidas.filter(saidas.ano == ano_int, saidas.mes > 6, saidas.capcus.like("CUSTEIO")).all()
        saidas_lista_capital = query_saidas.filter(saidas.ano == ano_int, saidas.mes > 6, saidas.capcus.like("CAPITAL")).all()
        total_saidas_custeio = sum(s.valor1 for s in saidas_lista_custeio) / 100
        total_saidas_capital = sum(s.valor1 for s in saidas_lista_capital) / 100

        saldo_final_custeio = total_receita_custeio - total_saidas_custeio
        saldo_final_capital = total_receita_capital - total_saidas_capital
        
        todas_saidas = saidas_lista_custeio + saidas_lista_capital
        
    # --- Lógica para o 1º SEMESTRE (periodo == 1) ---
    elif periodo == 1:
        # Lógica similar a periodo 3, mas filtrando por mes <= 6
        query_entradas_anterior_custeio = query_entradas.filter(entradas.ano <= ano_anterior, entradas.capcus.like("CUSTEIO"))
        query_saidas_anterior_custeio = query_saidas.filter(saidas.ano <= ano_anterior, saidas.capcus.like("CUSTEIO"))
        query_entradas_anterior_capital = query_entradas.filter(entradas.ano <= ano_anterior, entradas.capcus.like("CAPITAL"))
        query_saidas_anterior_capital = query_saidas.filter(saidas.ano <= ano_anterior, saidas.capcus.like("CAPITAL"))

        saldo_anterior_custeio = (sum(e.valor for e in query_entradas_anterior_custeio.all()) - sum(s.valor1 for s in query_saidas_anterior_custeio.all())) / 100
        saldo_anterior_capital = (sum(e.valor for e in query_entradas_anterior_capital.all()) - sum(s.valor1 for s in query_saidas_anterior_capital.all())) / 100

        entradas_semestre_custeio = query_entradas.filter(entradas.ano == ano_int, entradas.mes <= 6, entradas.capcus.like("CUSTEIO")).all()
        entradas_semestre_capital = query_entradas.filter(entradas.ano == ano_int, entradas.mes <= 6, entradas.capcus.like("CAPITAL")).all()
        
        total_entradas_proprio_custeio = sum(e.valor for e in entradas_semestre_custeio if e.comentario == "RECURSOS PRÓPRIOS") / 100
        total_entradas_proprio_capital = sum(e.valor for e in entradas_semestre_capital if e.comentario == "RECURSOS PRÓPRIOS") / 100
        total_entradas_rendimento_custeio = sum(e.valor for e in entradas_semestre_custeio if e.comentario == "RENDIMENTOS") / 100
        total_entradas_rendimento_capital = sum(e.valor for e in entradas_semestre_capital if e.comentario == "RENDIMENTOS") / 100
        total_entradas_recebido_custeio = (sum(e.valor for e in entradas_semestre_custeio) / 100) - total_entradas_proprio_custeio - total_entradas_rendimento_custeio
        total_entradas_recebido_capital = (sum(e.valor for e in entradas_semestre_capital) / 100) - total_entradas_proprio_capital - total_entradas_rendimento_capital

        total_receita_custeio = saldo_anterior_custeio + total_entradas_recebido_custeio + total_entradas_proprio_custeio + total_entradas_rendimento_custeio
        total_receita_capital = saldo_anterior_capital + total_entradas_recebido_capital + total_entradas_proprio_capital + total_entradas_rendimento_capital
        
        saidas_lista_custeio = query_saidas.filter(saidas.ano == ano_int, saidas.mes <= 6, saidas.capcus.like("CUSTEIO")).all()
        saidas_lista_capital = query_saidas.filter(saidas.ano == ano_int, saidas.mes <= 6, saidas.capcus.like("CAPITAL")).all()
        total_saidas_custeio = sum(s.valor1 for s in saidas_lista_custeio) / 100
        total_saidas_capital = sum(s.valor1 for s in saidas_lista_capital) / 100

        saldo_final_custeio = total_receita_custeio - total_saidas_custeio
        saldo_final_capital = total_receita_capital - total_saidas_capital

        todas_saidas = saidas_lista_custeio + saidas_lista_capital
    else:
        return "Período inválido selecionado.", 400

    # --- INÍCIO DA LÓGICA DE SEPARAÇÃO E ORDENAÇÃO (COMUM A TODOS OS PERÍODOS) ---

    # 1. Separar as saídas em duas listas: principal e bancos
    fornecedores_banco = ["CAIXA", "Banco do Brasil"]
    saidas_principais = [s for s in todas_saidas if s.fornecedor1 not in fornecedores_banco]
    saidas_bancos = [s for s in todas_saidas if s.fornecedor1 in fornecedores_banco]

    # 2. Ordenar ambas as listas pela data de pagamento
    key_ordenacao = lambda x: datetime.strptime(x.data, '%d/%m/%Y')
    saidas_principais_ordenadas = sorted(saidas_principais, key=key_ordenacao)
    saidas_bancos_ordenadas = sorted(saidas_bancos, key=key_ordenacao)

    # 3. Calcular as somas para as despesas bancárias separadamente
    total_despesas_caixa = sum(s.valor1 for s in saidas_bancos if s.fornecedor1 == "CAIXA") / 100
    total_despesas_bb = sum(s.valor1 for s in saidas_bancos if s.fornecedor1 == "Banco do Brasil") / 100

    # --- FIM DA LÓGICA DE SEPARAÇÃO ---


    # --- MONTAGEM DO HTML DE RESPOSTA ---

    def format_brl(value):
        """Formata um valor float para a moeda brasileira (R$ 1.234,56)."""
        return f"R$ {value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

    # Tabela de Resumo Financeiro
    nova_tabela_html = f"""
    <h2>Resumo Financeiro</h2>
    <table border='1'>
        <tr>
            <th>Descrição</th>
            <th>Custeio</th>
            <th>Capital</th>
        </tr>
        <tr><td>Saldo Anterior</td><td>{format_brl(saldo_anterior_custeio)}</td><td>{format_brl(saldo_anterior_capital)}</td></tr>
        <tr><td>Recursos Recebidos</td><td>{format_brl(total_entradas_recebido_custeio)}</td><td>{format_brl(total_entradas_recebido_capital)}</td></tr>
        <tr><td>Recursos Próprios</td><td>{format_brl(total_entradas_proprio_custeio)}</td><td>{format_brl(total_entradas_proprio_capital)}</td></tr>
        <tr><td>Rendimentos</td><td>{format_brl(total_entradas_rendimento_custeio)}</td><td>{format_brl(total_entradas_rendimento_capital)}</td></tr>
        <tr><td><b>Receita Total</b></td><td><b>{format_brl(total_receita_custeio)}</b></td><td><b>{format_brl(total_receita_capital)}</b></td></tr>
        <tr><td>Despesas</td><td>{format_brl(total_saidas_custeio)}</td><td>{format_brl(total_saidas_capital)}</td></tr>
        <tr><td><b>Saldo Final</b></td><td><b>{format_brl(saldo_final_custeio)}</b></td><td><b>{format_brl(saldo_final_capital)}</b></td></tr>
    </table>
    """

    # Tabela de Compras Principais
    compras_html = "<h2>Compras Realizadas</h2>"
    if saidas_principais_ordenadas:
        compras_html += "<table border='1'><tr><th>#</th><th>Fornecedor</th><th>CNPJ</th><th>Descrição</th><th>Cap/Cus</th><th>SubFonte</th><th>Tipo Nota</th><th>Nº Nota</th><th>Data Nota</th><th>Nº Pag.</th><th>Data Pag.</th><th>Valor</th></tr>"
        for i, saida in enumerate(saidas_principais_ordenadas, start=1):
            datanota_fmt = datetime.strptime(saida.datanota, '%d/%m/%Y').strftime('%d/%m/%Y')
            datapag_fmt = datetime.strptime(saida.data, '%d/%m/%Y').strftime('%d/%m/%Y')
            compras_html += f"<tr><td>{i}</td><td>{saida.fornecedor1}</td><td>{saida.cnpj1}</td><td>{saida.descricao}</td><td>{saida.capcus}</td><td>{saida.subfonte}</td><td>{saida.tiponota}</td><td>{saida.numnota}</td><td>{datanota_fmt}</td><td>{saida.numpag}</td><td>{datapag_fmt}</td><td>{format_brl(saida.valor1 / 100)}</td></tr>"
        compras_html += "</table>"
    else:
        compras_html += "<p>Nenhuma compra realizada no período.</p>"


    # Tabela de Despesas Bancárias
    bancos_html = ""
    if saidas_bancos_ordenadas:
        bancos_html += "<h2>Despesas Bancárias</h2>"
        if total_despesas_caixa > 0:
            bancos_html += f"<h4>Soma Total (CAIXA): {format_brl(total_despesas_caixa)}</h4>"
        if total_despesas_bb > 0:
            bancos_html += f"<h4>Soma Total (Banco do Brasil): {format_brl(total_despesas_bb)}</h4>"
        bancos_html += "<table border='1'><tr><th>#</th><th>Fornecedor</th><th>CNPJ</th><th>Descrição</th><th>Cap/Cus</th><th>SubFonte</th><th>Tipo Nota</th><th>Nº Nota</th><th>Data Nota</th><th>Nº Pag.</th><th>Data Pag.</th><th>Valor</th></tr>"
        for i, saida in enumerate(saidas_bancos_ordenadas, start=1):
            datanota_fmt = datetime.strptime(saida.datanota, '%d/%m/%Y').strftime('%d/%m/%Y')
            datapag_fmt = datetime.strptime(saida.data, '%d/%m/%Y').strftime('%d/%m/%Y')
            bancos_html += f"<tr><td>{i}</td><td>{saida.fornecedor1}</td><td>{saida.cnpj1}</td><td>{saida.descricao}</td><td>{saida.capcus}</td><td>{saida.subfonte}</td><td>{saida.tiponota}</td><td>{saida.numnota}</td><td>{datanota_fmt}</td><td>{saida.numpag}</td><td>{datapag_fmt}</td><td>{format_brl(saida.valor1 / 100)}</td></tr>"
        bancos_html += "</table>"
    
    # Concatenar o HTML final
    html_final = nova_tabela_html + compras_html + bancos_html

    return render_template_string(html_final)


#CHAMADA GERAL
@app.route('/', methods=['GET'])
def index():
    return redirect(url_for('homepage')) 


def open_browser():
     webbrowser.open_new('http://127.0.0.1:5000/')

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    Timer(1, open_browser).start() 
    app.run(debug=True)
 
